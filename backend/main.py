import os
import sys
import logging
from pathlib import Path

# Setup logging to file
log_dir = Path(__file__).parent
log_file = log_dir / "backend.log"

logging.basicConfig(
    level=logging.DEBUG,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[logging.FileHandler(log_file), logging.StreamHandler(sys.stdout)],
)
logger = logging.getLogger(__name__)

# Load environment variables from .env file
env_path = Path(__file__).parent.parent / ".env"
if env_path.exists():
    from dotenv import load_dotenv

    load_dotenv(dotenv_path=env_path)

from fastapi import FastAPI, HTTPException, Depends, BackgroundTasks, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from typing import Optional, List, Dict, Any
from pydantic import BaseModel, HttpUrl
from datetime import datetime
import httpx
import asyncio
import json as _json

from models import (
    init_db,
    get_db,
    Job,
    JobApplication,
    BaseResume,
    BaseCoverLetter,
    GeneratedResume,
    GeneratedCoverLetter,
    ArtifactScore,
    Subscription,
    UsageEvent,
    SessionLocal,
    User,
    generate_public_job_id,
)
from job_parser import JobParser
from agents import agent_service
from auth import get_current_user
from ssrf import is_url_safe
import billing

logger = logging.getLogger(__name__)

app = FastAPI(title="Joblign API")

# CORS
CORS_ORIGINS = os.getenv("CORS_ORIGIN", "http://localhost:8765,https://joblign.ronning.systems").split(",")
app.add_middleware(
    CORSMiddleware,
    allow_origins=CORS_ORIGINS,
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE", "PATCH"],
    allow_headers=["Authorization", "Content-Type"],
)


# Initialize database on startup
@app.on_event("startup")
async def startup_event():
    init_db()


# Auth endpoints
@app.get("/api/auth/me")
async def get_auth_me(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Return current authenticated user's profile + plan + usage"""
    plan = billing.get_effective_plan(current_user)
    cap = billing.plan_cap(plan)
    used = billing.get_usage_this_month(current_user, db)
    return {
        "id": current_user.id,
        "auth0_id": current_user.auth0_id,
        "email": current_user.email,
        "name": current_user.name,
        "avatar_url": current_user.avatar_url,
        "created_at": current_user.created_at.isoformat() if current_user.created_at else None,
        "last_login": current_user.last_login.isoformat() if current_user.last_login else None,
        "plan": plan,
        "plan_grandfathered": bool(getattr(current_user, "plan_grandfathered", False)),
        "billing": {
            "enabled": not billing.BILLING_DISABLED,
            "plan": plan,
            "cap": cap,
            "used": used,
            "remaining": max(0, cap - used),
        },
    }


# Pydantic models
class JobCreateInput(BaseModel):
    url: Optional[str] = None
    job_text: Optional[str] = None

    class Config:
        json_schema_extra = {
            "example": {"url": "https://example.com/job/123", "job_text": ""}
        }


class JobUpdate(BaseModel):
    company: Optional[str] = None
    position: Optional[str] = None
    location: Optional[str] = None
    salary: Optional[str] = None
    remote: Optional[str] = None
    job_url: Optional[str] = None
    job_description_raw: Optional[str] = None
    job_description_parsed: Optional[str] = None
    notes: Optional[str] = None
    applied_date: Optional[datetime] = None


class ApplicationUpdate(BaseModel):
    stage: Optional[str] = None
    applied_date: Optional[datetime] = None
    response_received: Optional[bool] = None
    notes: Optional[str] = None


class BaseResumeCreate(BaseModel):
    name: str
    resume_type: str  # 'example' or 'template'
    content: str  # Base64 encoded file content
    atoms: Optional[list] = None  # only for template type
    docx_base64: Optional[str] = None  # only for template type


class JobResponse(BaseModel):
    id: int
    public_job_id: Optional[str] = None
    company: str
    position: str
    location: Optional[str]
    salary: Optional[str]
    pay_range_min: Optional[int] = None
    pay_range_max: Optional[int] = None
    pay_currency: Optional[str] = "USD"
    pay_period: Optional[str] = None
    application_deadline: Optional[datetime] = None
    remote: Optional[str]
    job_url: Optional[str]
    stage: str
    applied_date: Optional[datetime]
    response_received: bool
    notes: Optional[str]
    keywords: Optional[list]
    generated_resume: Optional[str]
    resume_revisions: Optional[list] = []
    structured_content: Optional[dict] = None
    atoms_snapshot: Optional[list] = None
    has_cover_letter: bool = False
    created_at: datetime

    class Config:
        from_attributes = True


class JobDetailResponse(JobResponse):
    job_description_raw: Optional[str]
    job_description_parsed: Optional[str]
    requirements: Optional[dict]
    responsibilities: Optional[list]
    required_credentials: Optional[list]
    history: Optional[list]
    generated_resume: Optional[str]
    cover_letter: Optional[str] = None
    cover_letter_revisions: Optional[list] = []




@app.post("/api/jobs", response_model=JobDetailResponse)
async def create_job(
    input_data: JobCreateInput,
    background_tasks: BackgroundTasks,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """
    Create a new job from job description text.
    The job text is parsed using Ollama to extract structured data.
    The URL is stored for reference only (not fetched).
    """
    parser = JobParser()

    if not input_data.job_text:
        raise HTTPException(status_code=400, detail="Job description text is required")

    # Parse from plain text using Ollama
    try:
        job_data = await parser.parse_from_text(input_data.job_text)
    except Exception as e:
        raise HTTPException(
            status_code=500, detail=f"Failed to parse job text: {str(e)}"
        )

    # Extract structured pay range + deadline from parsed data (if present)
    pay_range = job_data.get("pay_range") or {}
    deadline_str = job_data.get("application_deadline")
    application_deadline = None
    if deadline_str:
        try:
            application_deadline = datetime.fromisoformat(deadline_str)
        except (ValueError, TypeError):
            application_deadline = None

    # Create Job record
    job = Job(
        public_job_id=_generate_unique_public_job_id(db),
        user_id=current_user.id,
        company=job_data.get("company", "Unknown"),
        position=job_data.get("position", "Unknown"),
        location=job_data.get("location"),
        salary=job_data.get("salary"),
        pay_range_min=pay_range.get("min"),
        pay_range_max=pay_range.get("max"),
        pay_currency=pay_range.get("currency", "USD"),
        pay_period=pay_range.get("period"),
        application_deadline=application_deadline,
        remote=job_data.get("remote"),
        job_url=input_data.url or job_data.get("url"),
        job_description_raw=job_data.get("raw_text", input_data.job_text),
        job_description_parsed=job_data.get("description"),
        requirements=job_data.get("requirements", {}),
        responsibilities=job_data.get("responsibilities", []),
        keywords=job_data.get("keywords", []),
        required_credentials=job_data.get("credentials", []),
        source_type="url" if input_data.url else "text",
        source_url=input_data.url,
    )

    db.add(job)
    db.commit()
    db.refresh(job)

    # Create initial application record with "saved" stage
    application = JobApplication(
        job_id=job.id,
        user_id=current_user.id,
        stage="saved",
        history=[
            {
                "date": datetime.utcnow().isoformat(),
                "action": "job_added",
                "notes": "Job added",
            }
        ],
    )
    db.add(application)
    db.commit()

    return format_job_response(job, application, db)


@app.get("/api/jobs", response_model=List[JobResponse])
def list_jobs(
    stage: Optional[str] = None,
    search: Optional[str] = None,
    sort: Optional[str] = None,
    order: Optional[str] = "desc",
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """List all jobs for the current user with optional filtering + sorting.

    sort ∈ {company, position, location, stage, applied_date, deadline, pay, created}.
    order ∈ {asc, desc} (default desc).
    """
    query = db.query(Job, JobApplication).join(
        JobApplication, Job.id == JobApplication.job_id
    ).filter(Job.user_id == current_user.id)

    if stage:
        query = query.filter(JobApplication.stage == stage)

    if search:
        search_filter = f"%{search}%"
        query = query.filter(
            (Job.company.ilike(search_filter))
            | (Job.position.ilike(search_filter))
            | (Job.location.ilike(search_filter))
        )

    # Sorting
    sort_map = {
        "company": Job.company,
        "position": Job.position,
        "location": Job.location,
        "stage": JobApplication.stage,
        "applied_date": JobApplication.applied_date,
        "deadline": Job.application_deadline,
        "pay": Job.pay_range_max,
        "created": Job.created_at,
    }
    sort_col = sort_map.get(sort, Job.created_at)
    if order and order.lower() == "asc":
        results = query.order_by(sort_col.asc()).all()
    else:
        results = query.order_by(sort_col.desc()).all()

    return [format_job_response(job, app, db) for job, app in results]


@app.get("/api/jobs/{job_id}", response_model=JobDetailResponse)
def get_job(job_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """Get detailed job information"""
    result = (
        db.query(Job, JobApplication)
        .join(JobApplication, Job.id == JobApplication.job_id)
        .filter(Job.id == job_id, Job.user_id == current_user.id)
        .first()
    )

    if not result:
        raise HTTPException(status_code=404, detail="Job not found")

    job, application = result
    return format_job_response(job, application, db)


# Debug endpoint to check resume revisions in DB
@app.get("/api/debug/resume/{job_id}")
def debug_resume(job_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """Debug: Check what's actually stored in the GeneratedResume table"""
    resumes = db.query(GeneratedResume).filter(GeneratedResume.job_id == job_id).all()
    return {
        "job_id": job_id,
        "resume_count": len(resumes),
        "resumes": [
            {
                "id": r.id,
                "job_id": r.job_id,
                "current_content_length": len(r.current_content) if r.current_content else 0,
                "revisions": r.revisions,
                "created_at": r.created_at.isoformat() if r.created_at else None,
                "updated_at": r.updated_at.isoformat() if r.updated_at else None,
            }
            for r in resumes
        ]
    }


@app.put("/api/jobs/{job_id}", response_model=JobDetailResponse)
def update_job(job_id: int, update: JobUpdate, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """Update job details"""
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == current_user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    update_dict = update.dict(exclude_unset=True)

    # Extract applied_date to update on application instead of job
    applied_date = update_dict.pop("applied_date", None)

    for field, value in update_dict.items():
        if hasattr(job, field):
            setattr(job, field, value)

    # Update applied_date on the application if provided
    if applied_date:
        application = (
            db.query(JobApplication).filter(JobApplication.job_id == job_id).first()
        )
        if application:
            application.applied_date = applied_date

    job.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(job)

    application = (
        db.query(JobApplication).filter(JobApplication.job_id == job_id).first()
    )
    return format_job_response(job, application, db)


@app.delete("/api/jobs/{job_id}")
def delete_job(job_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """Delete a job and its application records"""
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == current_user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    db.delete(job)
    db.commit()

    return {"message": "Job deleted successfully"}


@app.post("/api/resumes/base")
def create_base_resume(resume: BaseResumeCreate, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """Create or update a base resume (example or template)"""
    # For template type, delete any existing template for this user first.
    # This enforces the "single template per user, replacement overwrites"
    # behavior — the frontend must confirm with the user before calling.
    replaced_existing = False
    if resume.resume_type == "template":
        existing = db.query(BaseResume).filter(
            BaseResume.resume_type == "template",
            BaseResume.user_id == current_user.id,
        ).first()
        if existing is not None:
            replaced_existing = True
            db.delete(existing)
            db.flush()

    # Create new resume record
    base_resume = BaseResume(
        name=resume.name,
        resume_type=resume.resume_type,
        content=resume.content,
        source="upload",
        user_id=current_user.id,
    )
    if resume.resume_type == "template":
        # Persist the parsed atoms + raw DOCX so the composer can reproduce
        # the template's exact styling later.
        if resume.atoms is not None:
            base_resume.atoms_json = resume.atoms
        if resume.docx_base64 is not None:
            base_resume.docx_base64 = resume.docx_base64
    db.add(base_resume)
    db.commit()
    db.refresh(base_resume)

    msg = f"{resume.resume_type.capitalize()} resume saved successfully"
    if replaced_existing:
        msg = "Existing template replaced. " + msg
    return {
        "id": base_resume.id,
        "name": base_resume.name,
        "resume_type": base_resume.resume_type,
        "replaced_existing": replaced_existing,
        "atoms_count": len(resume.atoms) if resume.atoms else 0,
        "message": msg,
    }


@app.post("/api/template/parse")
def parse_template_endpoint(request: dict, current_user: User = Depends(get_current_user)):
    """Parse a template DOCX (without saving) and return its atoms + a preview.

    The frontend calls this when the user uploads a new template, to:
      1) Show the user the detected atoms / structure BEFORE they confirm
         replacing their current template.
      2) Get the docx_base64 back so it can pass it to /api/resumes/base.
    """
    from template_engine import parse_template_from_b64
    content_b64 = request.get("content_b64", "")
    if not content_b64:
        raise HTTPException(status_code=400, detail="content_b64 is required")
    try:
        result = parse_template_from_b64(content_b64)
    except Exception as e:
        logger.error(f"Template parse failed: {e}", exc_info=True)
        raise HTTPException(status_code=400, detail=f"Failed to parse DOCX: {e}")

    # Return everything except the raw docx bytes (frontend already has them);
    # include docx_base64 so the frontend can pass it straight to the save call.
    return {
        "atoms": result["atoms"],
        "docx_base64": result["docx_base64"],
        "page_setup": result["page_setup"],
        "canonical_atoms_present": result["canonical_atoms_present"],
        "warnings": result["warnings"],
    }


@app.get("/api/template")
def get_current_template(db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """Return the current template's atoms (and its raw DOCX) so the frontend
    can re-render the structured editor after page reloads."""
    tmpl = db.query(BaseResume).filter(
        BaseResume.resume_type == "template",
        BaseResume.user_id == current_user.id,
    ).first()
    if not tmpl:
        return {"has_template": False}
    return {
        "has_template": True,
        "id": tmpl.id,
        "name": tmpl.name,
        "atoms": tmpl.atoms_json or [],
        "docx_base64": tmpl.docx_base64,
    }


class ComposeDocxRequest(BaseModel):
    structured_content: Dict[str, Any]
    template_id: Optional[int] = None  # not strictly needed — we use the
    # single per-user template — but allows future multi-template support


@app.post("/api/resumes/compose-docx")
def compose_docx_endpoint(
    request: ComposeDocxRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Compose a DOCX from a structured content tree using the user's template.

    Returns the DOCX bytes as a base64-encoded string along with a filename
    suggestion. Frontend decodes + triggers a download.
    """
    from template_engine import compose_docx as te_compose_docx

    tmpl = db.query(BaseResume).filter(
        BaseResume.resume_type == "template",
        BaseResume.user_id == current_user.id,
    ).first()
    if not tmpl or not tmpl.docx_base64 or not tmpl.atoms_json:
        raise HTTPException(
            status_code=400,
            detail="No template configured. Upload a template in Resume Settings first.",
        )

    import base64 as _b64
    template_bytes = _b64.b64decode(tmpl.docx_base64)
    try:
        out_bytes = te_compose_docx(template_bytes, tmpl.atoms_json, request.structured_content)
    except Exception as e:
        logger.error(f"compose_docx failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to compose DOCX: {e}")

    return {
        "docx_base64": _b64.b64encode(out_bytes).decode("ascii"),
        "filename_suggestion": "resume.docx",
        "size_bytes": len(out_bytes),
    }


class StructuredEditRequest(BaseModel):
    structured_content: Dict[str, Any]
    current_content: Optional[str] = None


@app.put("/api/resumes/structured/{job_id}")
def update_structured_resume(
    job_id: int,
    request: StructuredEditRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Persist manual edits to the latest revision's structured content.

    Updates the latest GeneratedResume's structured_content + current_content
    fields. Does NOT create a new revision (the user is editing, not
    regenerating). The next regenerate will overwrite.
    """
    from sqlalchemy.orm.attributes import flag_modified

    job = db.query(Job).filter(Job.id == job_id, Job.user_id == current_user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    resume = db.query(GeneratedResume).filter(
        GeneratedResume.job_id == job_id
    ).order_by(GeneratedResume.updated_at.desc()).first()
    if not resume:
        raise HTTPException(status_code=404, detail="No generated resume to edit")

    resume.structured_content = request.structured_content
    flag_modified(resume, "structured_content")
    if request.current_content is not None:
        resume.current_content = request.current_content
    resume.updated_at = datetime.utcnow()
    # Also update the latest revision in the JSON list to keep them in sync
    if resume.revisions:
        revs = list(resume.revisions or [])
        if revs:
            revs[-1]["structured_content"] = request.structured_content
            if request.current_content is not None:
                revs[-1]["content"] = request.current_content
            revs[-1]["edited_at"] = datetime.utcnow().isoformat()
            resume.revisions = list(revs)
            flag_modified(resume, "revisions")

    db.commit()
    db.refresh(resume)

    return {
        "ok": True,
        "job_id": job_id,
        "resume_id": resume.id,
    }


@app.get("/api/resumes/base")
def list_base_resumes(resume_type: Optional[str] = None, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """List all base resumes for the current user, optionally filtered by type"""
    query = db.query(BaseResume).filter(BaseResume.user_id == current_user.id)
    if resume_type:
        query = query.filter(BaseResume.resume_type == resume_type)

    resumes = query.all()
    out = []
    for r in resumes:
        item = {
            "id": r.id,
            "name": r.name,
            "resume_type": r.resume_type,
            "created_at": r.created_at.isoformat() if r.created_at else None,
        }
        if r.resume_type == "template":
            item["atoms_count"] = len(r.atoms_json) if r.atoms_json else 0
            item["has_docx"] = bool(r.docx_base64)
        out.append(item)
    return out


@app.get("/api/resumes/base/{resume_id}")
def get_base_resume(resume_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """Get a single base resume with extracted text content"""
    resume = db.query(BaseResume).filter(
        BaseResume.id == resume_id,
        BaseResume.user_id == current_user.id
    ).first()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    # Extract text from the stored content
    extracted_text = ""
    if resume.content:
        try:
            from agents import extract_text_from_file
            extracted_text = extract_text_from_file(resume.content, resume.name)
        except Exception as e:
            logger.error(f"Failed to extract text from resume {resume_id}: {e}")
            extracted_text = ""

    response = {
        "id": resume.id,
        "name": resume.name,
        "resume_type": resume.resume_type,
        "created_at": resume.created_at.isoformat() if resume.created_at else None,
        "extracted_text": extracted_text,
    }
    if resume.resume_type == "template":
        response["atoms"] = resume.atoms_json or []
        response["has_docx"] = bool(resume.docx_base64)
    return response


@app.delete("/api/resumes/base/{resume_id}")
def delete_base_resume(resume_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """Delete a base resume"""
    resume = db.query(BaseResume).filter(BaseResume.id == resume_id, BaseResume.user_id == current_user.id).first()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    db.delete(resume)
    db.commit()

    return {"message": "Resume deleted successfully"}


@app.patch("/api/jobs/{job_id}/stage", response_model=JobDetailResponse)
def update_stage(job_id: int, update: ApplicationUpdate, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """Update the application stage and related fields"""
    # First verify the job belongs to this user
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == current_user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    application = (
        db.query(JobApplication).filter(JobApplication.job_id == job_id).first()
    )
    if not application:
        raise HTTPException(status_code=404, detail="Application not found")

    # Add history entry if stage changed
    if update.stage and update.stage != application.stage:
        history_entry = {
            "date": datetime.utcnow().isoformat(),
            "action": f"stage_changed",
            "from": application.stage,
            "to": update.stage,
            "notes": update.notes or f"Stage changed to {update.stage}",
        }
        application.history = (application.history or []) + [history_entry]
        application.stage = update.stage

    # Update other fields
    if update.applied_date is not None:
        application.applied_date = update.applied_date
    if update.response_received is not None:
        application.response_received = update.response_received
        if update.response_received and not application.response_date:
            application.response_date = datetime.utcnow()
    if update.notes is not None:
        application.notes = update.notes

    application.updated_at = datetime.utcnow()
    db.commit()
    db.refresh(application)

    job = db.query(Job).filter(Job.id == job_id).first()
    return format_job_response(job, application, db)


# In-memory tracking for background resume generation
_generation_status = {}  # {job_id: {"status": "processing"|"completed"|"error", "error": str, "version": int, "resume_id": int}}


def _set_generation_status(job_id: int, status: str, **kwargs):
    """Thread-safe helper to update the in-memory generation status."""
    _generation_status[job_id] = {
        "status": status,
        "updated_at": datetime.utcnow().isoformat(),
        **kwargs,
    }


def _progress(job_id: int, step: str, percent: int):
    """Post a progress update for a running generation."""
    _set_generation_status(job_id, "processing", step=step, percent=percent)
    logger.info(f"[generate-resume-progress] job_id={job_id} step={step} percent={percent}")


def _do_generate_resume(job_id: int, user_id: int, job_description: str, example_resumes: list, template, target_role: str, model_override: str, atoms: Optional[list] = None):
    """Run resume generation in a background thread, store result in DB and update status dict."""
    from sqlalchemy.orm.attributes import flag_modified
    import time
    overall_start = time.monotonic()
    try:
        _progress(job_id, "Preparing inputs", 5)
        # Run the async agent in a new event loop (NO DB session held during generation)
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            _progress(job_id, "Generating resume content", 25)
            gen_start = time.monotonic()
            resume_result = loop.run_until_complete(
                agent_service.generate_resume(
                    user_profile={},
                    job_description=job_description,
                    example_resumes=example_resumes,
                    template=template,
                    target_role=target_role,
                    model_override=model_override,
                    atoms=atoms,
                )
            )
            gen_ms = int((time.monotonic() - gen_start) * 1000)
            llm_ms = resume_result.get("llm_duration_ms")
            logger.info(f"[generate-resume-bg] job_id={job_id} agent path: {gen_ms}ms (LLM call: {llm_ms}ms)")
            _progress(job_id, "Processing generated content", 65)
        finally:
            loop.close()

        resume_content = resume_result.get("content", _json.dumps(resume_result))
        structured_content = resume_result.get("structured_content")  # may be None

        _progress(job_id, "Saving resume revision", 85)
        # Now open a DB session ONLY to save the result (short-lived)
        db = SessionLocal()
        try:
            model_label = resume_result.get("model_used") or model_override or "qwen3.5:cloud"

            existing_resume = db.query(GeneratedResume).filter(
                GeneratedResume.job_id == job_id,
                GeneratedResume.user_id == user_id,
            ).first()

            if existing_resume:
                # Coerce revisions into a plain list (JSON columns may round-trip as a custom type)
                revisions = list(existing_resume.revisions or [])
                # Self-heal version numbers — find the max version already present, then increment.
                # Falls back to length+1 if any revision is missing a version field.
                max_version = 0
                for r in revisions:
                    if isinstance(r, dict) and isinstance(r.get("version"), int):
                        if r["version"] > max_version:
                            max_version = r["version"]
                next_version = max_version + 1 if max_version else len(revisions) + 1

                logger.info(
                    f"[generate-resume-bg] Regenerating job {job_id}, existing revisions: {len(revisions)}, next version: {next_version}"
                )

                new_revision = {
                    "version": next_version,
                    "content": resume_content,
                    "model": model_label,
                    "feedback": None,
                    "timestamp": datetime.utcnow().isoformat(),
                }
                if structured_content is not None:
                    new_revision["structured_content"] = structured_content
                if atoms is not None:
                    new_revision["atoms_used"] = [a.get("id") for a in atoms]
                revisions.append(new_revision)

                # Assign a brand-new list (not the same reference we appended to) so SQLAlchemy
                # is guaranteed to see the change. flag_modified is the belt-and-suspenders for
                # JSON columns on backends where in-place mutation may not trigger a write.
                existing_resume.current_content = resume_content
                existing_resume.revisions = list(revisions)
                if structured_content is not None:
                    existing_resume.structured_content = structured_content
                    flag_modified(existing_resume, "structured_content")
                if atoms is not None:
                    existing_resume.atoms_snapshot = [a.get("id") for a in atoms]
                    flag_modified(existing_resume, "atoms_snapshot")
                flag_modified(existing_resume, "revisions")
                existing_resume.updated_at = datetime.utcnow()
                generated_resume = existing_resume

                logger.info(
                    f"[generate-resume-bg] Appended revision v{next_version} with model {model_label}, total revisions now: {len(revisions)}"
                )
            else:
                initial_revision = {
                    "version": 1,
                    "content": resume_content,
                    "model": model_label,
                    "feedback": None,
                    "timestamp": datetime.utcnow().isoformat(),
                }
                if structured_content is not None:
                    initial_revision["structured_content"] = structured_content
                if atoms is not None:
                    initial_revision["atoms_used"] = [a.get("id") for a in atoms]
                generated_resume = GeneratedResume(
                    job_id=job_id,
                    user_id=user_id,
                    current_content=resume_content,
                    revisions=[initial_revision],
                )
                if structured_content is not None:
                    generated_resume.structured_content = structured_content
                if atoms is not None:
                    generated_resume.atoms_snapshot = [a.get("id") for a in atoms]
                db.add(generated_resume)

            job = db.query(Job).filter(Job.id == job_id).first()
            if job:
                job.updated_at = datetime.utcnow()

            # Billing: count this generation as a usage event. We insert it
            # right after the resume is committed so we only count successes.
            # The cap was already enforced at the sync /generate-resume
            # endpoint, so this is the second leg of the gate.
            try:
                if not billing.BILLING_DISABLED:
                    usage_user = db.query(User).filter(User.id == user_id).first()
                    if usage_user:
                        event = UsageEvent(
                            user_id=user_id,
                            event_type="resume_generation",
                            job_id=job_id,
                            plan_at_event=billing.get_effective_plan(usage_user),
                        )
                        db.add(event)
            except Exception as e:
                # Don't fail the whole generation if the usage insert hiccups;
                # the gate at the endpoint will catch next time anyway.
                logger.warning(f"[generate-resume-bg] usage insert failed: {e}")

            db.commit()
            db.refresh(generated_resume)

            # Quality flag: if structured_content is mostly empty, the LLM
            # probably produced valid JSON but with blank fields (common with
            # small/cheap models). Surface that to the frontend so it can
            # warn the user to switch models.
            empty_ratio = 0.0
            if structured_content is not None and isinstance(structured_content, dict):
                atoms_list = structured_content.get("atoms") or []
                if atoms_list:
                    empty = 0
                    for a in atoms_list:
                        if not isinstance(a, dict):
                            continue
                        has_content = False
                        if isinstance(a.get("text"), str) and a["text"].strip():
                            has_content = True
                        if not has_content and isinstance(a.get("segments"), list):
                            has_content = any(
                                isinstance(s, dict) and (s.get("text") or "").strip()
                                for s in a["segments"]
                            )
                        if not has_content:
                            empty += 1
                    empty_ratio = empty / len(atoms_list)

            total_ms = int((time.monotonic() - overall_start) * 1000)
            _set_generation_status(
                job_id,
                "completed",
                resume_id=generated_resume.id,
                version=len(generated_resume.revisions) if generated_resume.revisions else 1,
                mode=resume_result.get("mode", "plain"),
                structured_empty_ratio=round(empty_ratio, 3),
                step="Done",
                percent=100,
            )
            logger.info(
                f"[generate-resume-bg] Completed job {job_id}: "
                f"mode={resume_result.get('mode')}, "
                f"structured_content atoms={len(structured_content.get('atoms', [])) if structured_content else 0}, "
                f"empty_ratio={round(empty_ratio, 3)}, "
                f"total_time_ms={total_ms}, "
                f"llm_time_ms={resume_result.get('llm_duration_ms', 'n/a')}"
            )
        finally:
            db.close()
    except Exception as e:
        logger.error(f"[generate-resume-bg] Error for job {job_id}: {e}", exc_info=True)
        _set_generation_status(job_id, "error", error=str(e), step="Error", percent=0)


@app.post("/api/jobs/{job_id}/generate-resume")
async def generate_job_resume(
    job_id: int, resume_request: dict, background_tasks: BackgroundTasks, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)
):
    """Generate a resume for a specific job using example resumes and template.
    Runs generation in background; client polls /generate-resume/status/{job_id}."""
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == current_user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    # Billing: gate the request before we do any work. We check (don't
    # increment) here; the actual UsageEvent row is written in
    # _do_generate_resume only on success.
    plan = billing.get_effective_plan(current_user)
    cap = billing.plan_cap(plan)
    used = billing.get_usage_this_month(current_user, db)
    if used >= cap and not billing.BILLING_DISABLED:
        # 402 Payment Required — frontend recognizes this and opens paywall.
        raise HTTPException(
            status_code=402,
            detail={
                "code": "usage_cap_reached",
                "message": f"You've used all {cap} {billing.PLANS[plan]['name']} generations this month. Upgrade to Pro for more.",
                "plan": plan,
                "cap": cap,
                "used": used,
            },
        )

    job_description = job.job_description_parsed or job.job_description_raw

    # Debug: Log what's received
    logger.info(f"[generate-resume] Received request body: {resume_request}")
    logger.info(
        f"[generate-resume] example_resumes: {resume_request.get('example_resumes')}"
    )
    logger.info(f"[generate-resume] template: {resume_request.get('template')}")

    # Get example_resumes and template from request body first
    example_resumes = resume_request.get("example_resumes", [])
    template = resume_request.get("template")

    logger.info(
        f"[generate-resume] After get - example_resumes: {example_resumes}, template: {template}"
    )

    # Fall back to database if not provided in request, or enrich with content if missing
    if not example_resumes:
        example_resumes_db = (
            db.query(BaseResume).filter(BaseResume.resume_type == "example", BaseResume.user_id == current_user.id).all()
        )
        example_resumes = [
            {"name": r.name, "content": r.content} for r in example_resumes_db
        ]
        logger.info(
            f"[generate-resume] Fell back to DB - example_resumes: {len(example_resumes)} found"
        )
    else:
        # Frontend sends metadata only (id, name) without content — enrich from DB
        for er in example_resumes:
            if not er.get("content"):
                resume_id = er.get("id")
                if resume_id:
                    db_resume = db.query(BaseResume).filter(
                        BaseResume.id == resume_id,
                        BaseResume.user_id == current_user.id
                    ).first()
                    if db_resume:
                        er["content"] = db_resume.content
                        logger.info(f"[generate-resume] Enriched example resume {er.get('name')} with DB content")
                else:
                    # No id — try to find by name
                    db_resume = db.query(BaseResume).filter(
                        BaseResume.resume_type == "example",
                        BaseResume.user_id == current_user.id,
                        BaseResume.name == er.get("name")
                    ).first()
                    if db_resume:
                        er["content"] = db_resume.content
                        logger.info(f"[generate-resume] Enriched example resume {er.get('name')} with DB content (by name)")

    if not template:
        template_db = (
            db.query(BaseResume).filter(BaseResume.resume_type == "template", BaseResume.user_id == current_user.id).first()
        )
        template = (
            {"name": template_db.name, "content": template_db.content}
            if template_db
            else None
        )
        logger.info(
            f"[generate-resume] Fell back to DB - template: {template_db.name if template_db else None}"
        )

    # Get optional model override from request
    model_override = resume_request.get("model")
    logger.info(f"[generate-resume] Model override: {model_override}")

    # Load the user's saved template atoms (if any). When present, the LLM
    # emits structured JSON; otherwise we fall back to plain-text mode.
    tmpl_row = db.query(BaseResume).filter(
        BaseResume.resume_type == "template",
        BaseResume.user_id == current_user.id,
    ).first()
    atoms = None
    if tmpl_row and tmpl_row.atoms_json:
        atoms = tmpl_row.atoms_json
        logger.info(
            f"[generate-resume] Template atoms loaded: {[a.get('id') for a in atoms]}"
        )
    else:
        logger.info(
            f"[generate-resume] No template atoms found in DB for user {current_user.id}"
        )

    # Resolve actual model: override wins, else env var
    actual_model = model_override or os.getenv("MODEL_GENERATION") or os.getenv("MODEL_AGENTS", "qwen3.5:cloud")
    logger.info(f"[generate-resume] Using model: {actual_model} (override={model_override})")

    # Launch background generation
    _set_generation_status(job_id, "processing", step="Queued", percent=0)
    background_tasks.add_task(
        _do_generate_resume,
        job_id=job_id,
        user_id=current_user.id,
        job_description=job_description,
        example_resumes=example_resumes,
        template=template,
        target_role=job.position,
        model_override=model_override,
        atoms=atoms,
    )

    return {"job_id": job_id, "status": "processing"}


@app.get("/api/jobs/{job_id}/generate-resume/status")
async def get_generate_resume_status(
    job_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)
):
    """Poll this endpoint to check background resume generation status."""
    # Verify job ownership
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == current_user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    status = _generation_status.get(job_id, {"status": "unknown"})

    # If completed, return the full result
    if status.get("status") == "completed":
        # Fetch the latest resume from DB
        latest = db.query(GeneratedResume).filter(
            GeneratedResume.job_id == job_id
        ).order_by(GeneratedResume.updated_at.desc()).first()

        if latest:
            sc = latest.structured_content
            logger.info(
                f"[status] job_id={job_id} status=completed, "
                f"structured_content={'present (' + str(len(sc.get('atoms', [])) if sc else 0) + ' atoms)'} if sc else 'None/Missing', "
                f"mode={status.get('mode')}, "
                f"empty_ratio={status.get('structured_empty_ratio', 0.0)}"
            )
            return {
                "status": "completed",
                "job_id": job_id,
                "resume": latest.current_content,
                "resume_id": latest.id,
                "version": len(latest.revisions) if latest.revisions else 1,
                "revisions": latest.revisions or [],
                "structured_content": latest.structured_content,
                "atoms_snapshot": latest.atoms_snapshot,
                "mode": status.get("mode", "plain"),
                "structured_empty_ratio": status.get("structured_empty_ratio", 0.0),
                "step": status.get("step", "Done"),
                "percent": status.get("percent", 100),
            }
        return {"status": "completed", "job_id": job_id, "resume": None, "step": "Done", "percent": 100}

    return {
        "status": status.get("status", "unknown"),
        "error": status.get("error"),
        "step": status.get("step"),
        "percent": status.get("percent"),
    }


@app.post("/api/jobs/{job_id}/revise-resume")
async def revise_job_resume(
    job_id: int,
    request: dict,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Revise a generated resume based on user feedback. Appends a new revision."""
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == current_user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    feedback = request.get("feedback", "")
    if not feedback:
        raise HTTPException(status_code=400, detail="Feedback is required")

    # Get existing resume
    existing_resume = db.query(GeneratedResume).filter(
        GeneratedResume.job_id == job_id
    ).first()
    if not existing_resume or not existing_resume.current_content:
        raise HTTPException(status_code=404, detail="No generated resume found. Generate one first.")

    # Get example resumes and template
    example_resumes = request.get("example_resumes", [])
    template = request.get("template", None)

    if not example_resumes:
        example_resumes_db = db.query(BaseResume).filter(
            BaseResume.resume_type == "example",
            BaseResume.user_id == current_user.id,
        ).all()
        example_resumes = [{"name": r.name, "content": r.content} for r in example_resumes_db]

    if not template:
        template_db = db.query(BaseResume).filter(
            BaseResume.resume_type == "template",
            BaseResume.user_id == current_user.id,
        ).first()
        template = {"name": template_db.name, "content": template_db.content} if template_db else None

    # Generate revised resume using agent service
    tmpl_row = db.query(BaseResume).filter(
        BaseResume.resume_type == "template",
        BaseResume.user_id == current_user.id,
    ).first()
    atoms_for_revise = tmpl_row.atoms_json if (tmpl_row and tmpl_row.atoms_json) else None
    model_override = request.get("model")
    resume_result = await agent_service.revise_resume(
        current_resume=existing_resume.current_content,
        feedback=feedback,
        job_description=job.job_description_parsed or job.job_description_raw,
        example_resumes=example_resumes,
        template=template,
        target_role=job.position,
        atoms=atoms_for_revise,
        current_structured=existing_resume.structured_content,
        model_override=model_override,
    )

    import json
    resume_content = resume_result.get("content", json.dumps(resume_result))

    # Append new revision with model info, robust against stale version numbers
    from sqlalchemy.orm.attributes import flag_modified
    revisions = list(existing_resume.revisions or [])
    max_version = 0
    for r in revisions:
        if isinstance(r, dict) and isinstance(r.get("version"), int):
            if r["version"] > max_version:
                max_version = r["version"]
    next_version = max_version + 1 if max_version else len(revisions) + 1
    model_used = resume_result.get("model_used") or os.getenv("MODEL_GENERATION") or os.getenv("MODEL_AGENTS", "qwen3.5:cloud")
    new_revision = {
        "version": next_version,
        "content": resume_content,
        "model": model_used,
        "feedback": feedback,
        "timestamp": datetime.utcnow().isoformat(),
    }
    revised_structured = resume_result.get("structured_content")
    if revised_structured is not None:
        new_revision["structured_content"] = revised_structured
    if atoms_for_revise is not None:
        new_revision["atoms_used"] = [a.get("id") for a in atoms_for_revise]
    revisions.append(new_revision)
    existing_resume.current_content = resume_content
    existing_resume.revisions = list(revisions)
    if revised_structured is not None:
        existing_resume.structured_content = revised_structured
        flag_modified(existing_resume, "structured_content")
    flag_modified(existing_resume, "revisions")
    existing_resume.updated_at = datetime.utcnow()
    job.updated_at = datetime.utcnow()
    db.commit()

    return {
        "job_id": job_id,
        "resume": resume_content,
        "resume_id": existing_resume.id,
        "version": next_version,
        "revisions": revisions,
        "structured_content": revised_structured,
        "mode": resume_result.get("mode", "plain"),
    }


@app.get("/api/health")
async def health_check():
    """Health check endpoint"""
    from auth import AUTH_DISABLED
    return {
        "status": "healthy",
        "service": "joblign-api",
        "auth_disabled": AUTH_DISABLED,
    }


@app.get("/api/health/ollama")
async def ollama_health_check():
    """Check if Ollama is accessible"""
    ollama_url = os.getenv("MODEL_ENDPOINT", "http://localhost:11434").rstrip("/")
    try:
        async with httpx.AsyncClient(timeout=10.0) as client:
            response = await client.get(f"{ollama_url}/api/tags")
            if response.status_code == 200:
                models = response.json().get("models", [])
                model_names = [m.get("name") for m in models]
                return {"status": "connected", "url": ollama_url, "models": model_names}
            else:
                return {
                    "status": "error",
                    "url": ollama_url,
                    "message": f"HTTP {response.status_code}",
                }
    except Exception as e:
        return {"status": "unreachable", "url": ollama_url, "message": str(e)}


# ----------------------------------------------------------------------------
# Billing endpoints
# ----------------------------------------------------------------------------

@app.get("/api/billing/config")
async def billing_config():
    """Public config for the SPA: publishable key, plan info, billing enabled."""
    return billing.get_public_config()


@app.get("/api/billing/status")
async def billing_status(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Return the current user's plan + current-month usage."""
    plan = billing.get_effective_plan(current_user)
    cap = billing.plan_cap(plan)
    used = billing.get_usage_this_month(current_user, db)
    return {
        "enabled": not billing.BILLING_DISABLED,
        "plan": plan,
        "plan_grandfathered": bool(getattr(current_user, "plan_grandfathered", False)),
        "cap": cap,
        "used": used,
        "remaining": max(0, cap - used),
        "has_stripe_customer": bool(current_user.stripe_customer_id),
    }


class CheckoutRequest(BaseModel):
    return_url: Optional[str] = None


@app.post("/api/billing/checkout")
async def billing_checkout(
    body: CheckoutRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Create a Stripe Checkout session for the Pro plan. Returns the URL."""
    if billing.BILLING_DISABLED:
        raise HTTPException(
            status_code=503,
            detail="Billing is not configured on this server. Set STRIPE_SECRET_KEY and STRIPE_PRICE_ID_PRO to enable.",
        )
    try:
        url = billing.create_checkout_session(current_user, db, return_url=body.return_url)
    except Exception as e:
        logger.error(f"[billing] create_checkout_session failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Could not create checkout session.")
    return {"url": url}


class PortalRequest(BaseModel):
    return_url: Optional[str] = None


@app.post("/api/billing/portal")
async def billing_portal(
    body: PortalRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Create a Stripe Customer Portal session. Returns the URL."""
    if billing.BILLING_DISABLED:
        raise HTTPException(status_code=503, detail="Billing is not configured on this server.")
    if not current_user.stripe_customer_id:
        raise HTTPException(
            status_code=400,
            detail="No Stripe customer for this user. Subscribe first to manage billing.",
        )
    try:
        url = billing.create_portal_session(current_user, db, return_url=body.return_url)
    except Exception as e:
        logger.error(f"[billing] create_portal_session failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="Could not create portal session.")
    return {"url": url}


@app.post("/api/billing/webhook")
async def billing_webhook(request: Request, db: Session = Depends(get_db)):
    """Receive Stripe webhooks. No auth (Stripe signs them).

    Configure in Stripe Dashboard: endpoint URL = <this route>, events:
    customer.subscription.created/updated/deleted,
    customer.subscription.trial_will_end,
    checkout.session.completed,
    invoice.payment_succeeded, invoice.payment_failed.
    """
    if billing.BILLING_DISABLED or not billing.STRIPE_WEBHOOK_SECRET:
        raise HTTPException(status_code=503, detail="Webhook not configured.")

    payload = await request.body()
    sig = request.headers.get("stripe-signature", "")
    try:
        event = stripe.Webhook.construct_event(
            payload, sig, billing.STRIPE_WEBHOOK_SECRET
        )
    except ValueError as e:
        logger.warning(f"[stripe-webhook] Invalid payload: {e}")
        raise HTTPException(status_code=400, detail="Invalid payload")
    except Exception as e:  # stripe.error.SignatureVerificationError
        logger.warning(f"[stripe-webhook] Invalid signature: {e}")
        raise HTTPException(status_code=400, detail="Invalid signature")

    try:
        status = billing.handle_webhook_event(event.to_dict() if hasattr(event, "to_dict") else dict(event), db)
    except Exception as e:
        logger.error(f"[stripe-webhook] handler error: {e}", exc_info=True)
        # 200 so Stripe doesn't retry on our bug; we've logged.
        return {"received": True, "error": str(e)}
    return {"received": True, "status": status}


@app.get("/api/stats")
def get_stats(db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """Get dashboard statistics for the current user"""
    from sqlalchemy import func

    stats = {"total": db.query(Job).filter(Job.user_id == current_user.id).count(), "by_stage": {}}

    # Count by stage (filtered by user)
    stage_counts = (
        db.query(JobApplication.stage, func.count(JobApplication.id))
        .join(Job, JobApplication.job_id == Job.id)
        .filter(Job.user_id == current_user.id)
        .group_by(JobApplication.stage)
        .all()
    )

    for stage, count in stage_counts:
        stats["by_stage"][stage] = count

    # Active = saved, applied, phone_screen, interview, executive_call, offered
    active_stages = ["saved", "applied", "phone_screen", "interview", "executive_call", "offered"]
    archived_stages = ["rejected", "withdrawn", "closed", "not_pursuing"]
    stats["active"] = sum(stats["by_stage"].get(s, 0) for s in active_stages)
    stats["interviews"] = stats["by_stage"].get("interview", 0) + stats["by_stage"].get(
        "executive_call", 0
    )
    stats["offers"] = stats["by_stage"].get("offered", 0)
    stats["archived"] = sum(stats["by_stage"].get(s, 0) for s in archived_stages)

    return stats


def _generate_unique_public_job_id(db: Session) -> str:
    """Generate a public_job_id that doesn't collide with any existing row."""
    for _ in range(20):
        candidate = generate_public_job_id()
        existing = db.query(Job).filter(Job.public_job_id == candidate).first()
        if not existing:
            return candidate
    # Astronomically unlikely — fall back with extra entropy
    return generate_public_job_id(length=12)


def format_job_response(
    job: Job, application: JobApplication, db: Session = None
) -> dict:
    """Format job and application into response"""
    # Ensure requirements has the correct structure
    requirements = job.requirements or {}
    if isinstance(requirements, dict):
        # Ensure must_have and nice_to_have keys exist
        if "must_have" not in requirements:
            requirements["must_have"] = []
        if "nice_to_have" not in requirements:
            requirements["nice_to_have"] = []
    else:
        # If requirements is not a dict, create empty structure
        requirements = {"must_have": [], "nice_to_have": []}

    # Get latest generated resume + cover letter
    generated_resume = None
    resume_revisions = []
    latest_resume = None
    cover_letter = None
    cover_letter_revisions = []
    if db:
        latest_resume = (
            db.query(GeneratedResume)
            .filter(GeneratedResume.job_id == job.id)
            .order_by(GeneratedResume.created_at.desc())
            .first()
        )
        if latest_resume:
            generated_resume = latest_resume.current_content
            resume_revisions = latest_resume.revisions or []

        latest_cover_letter = (
            db.query(GeneratedCoverLetter)
            .filter(GeneratedCoverLetter.job_id == job.id)
            .order_by(GeneratedCoverLetter.created_at.desc())
            .first()
        )
        if latest_cover_letter:
            cover_letter = latest_cover_letter.current_content
            cover_letter_revisions = latest_cover_letter.revisions or []

    return {
        "id": job.id,
        "public_job_id": job.public_job_id,
        "company": job.company,
        "position": job.position,
        "location": job.location,
        "salary": job.salary,
        "pay_range_min": job.pay_range_min,
        "pay_range_max": job.pay_range_max,
        "pay_currency": job.pay_currency,
        "pay_period": job.pay_period,
        "application_deadline": job.application_deadline,
        "remote": job.remote,
        "job_url": job.job_url,
        "stage": application.stage,
        "applied_date": application.applied_date,
        "response_received": application.response_received,
        "notes": application.notes,
        "keywords": job.keywords or [],
        "created_at": job.created_at,
        # Detail fields
        "job_description_raw": job.job_description_raw,
        "job_description_parsed": job.job_description_parsed,
        "requirements": requirements,
        "responsibilities": job.responsibilities or [],
        "required_credentials": job.required_credentials or [],
        "history": application.history or [],
        "generated_resume": generated_resume,
        "resume_revisions": resume_revisions,
        "structured_content": latest_resume.structured_content if latest_resume else None,
        "atoms_snapshot": latest_resume.atoms_snapshot if latest_resume else None,
        "has_cover_letter": cover_letter is not None,
        "cover_letter": cover_letter,
        "cover_letter_revisions": cover_letter_revisions,
    }


# ---- Base cover letter CRUD (parallel to base resumes) ----------------

class BaseCoverLetterCreate(BaseModel):
    name: str
    letter_type: str = "example"  # only 'example' in v1
    content: str  # Base64 encoded file content


@app.post("/api/cover-letters/base")
def create_base_cover_letter(letter: BaseCoverLetterCreate, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """Upload an example cover letter (voice/tone reference)."""
    base = BaseCoverLetter(
        name=letter.name,
        letter_type=letter.letter_type,
        content=letter.content,
        source="upload",
        user_id=current_user.id,
    )
    db.add(base)
    db.commit()
    db.refresh(base)
    return {"id": base.id, "name": base.name, "letter_type": base.letter_type, "created_at": base.created_at.isoformat()}


@app.get("/api/cover-letters/base")
def list_base_cover_letters(db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """List all example cover letters for the current user."""
    letters = db.query(BaseCoverLetter).filter(BaseCoverLetter.user_id == current_user.id).all()
    return [
        {
            "id": l.id,
            "name": l.name,
            "letter_type": l.letter_type,
            "created_at": l.created_at.isoformat() if l.created_at else None,
        }
        for l in letters
    ]


@app.delete("/api/cover-letters/base/{letter_id}")
def delete_base_cover_letter(letter_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """Delete an example cover letter."""
    letter = db.query(BaseCoverLetter).filter(
        BaseCoverLetter.id == letter_id,
        BaseCoverLetter.user_id == current_user.id,
    ).first()
    if not letter:
        raise HTTPException(status_code=404, detail="Cover letter not found")
    db.delete(letter)
    db.commit()
    return {"message": "Cover letter deleted"}


# ---- Cover letter generation + revision (background) -----------------

async def _do_generate_cover_letter(job_id: int, user_id: int, job_description: str, resume_content: Optional[str], example_cover_letters: list, target_role: Optional[str], company: Optional[str], position: Optional[str], model_override: Optional[str]):
    """Background cover letter generation (mirrors _do_generate_resume)."""
    db = SessionLocal()
    try:
        _progress(job_id, "Generating cover letter", 25)
        result = await agent_service.generate_cover_letter(
            job_description=job_description,
            resume_content=resume_content,
            example_cover_letters=example_cover_letters,
            target_role=target_role,
            company=company,
            position=position,
            model_override=model_override,
        )
        if "error" in result:
            _set_generation_status(job_id, "error", error=result["error"], step="Error", percent=0)
            return

        content = result.get("content", "")
        _progress(job_id, "Saving cover letter", 85)

        existing = db.query(GeneratedCoverLetter).filter(GeneratedCoverLetter.job_id == job_id).first()
        if existing:
            revisions = list(existing.revisions or [])
            max_version = max((r.get("version", 0) for r in revisions if isinstance(r, dict)), default=0)
            next_version = max_version + 1
            revisions.append({
                "version": next_version,
                "content": content,
                "model": result.get("model_used"),
                "feedback": None,
                "timestamp": datetime.utcnow().isoformat(),
            })
            existing.current_content = content
            existing.revisions = revisions
            from sqlalchemy.orm.attributes import flag_modified
            flag_modified(existing, "revisions")
            existing.updated_at = datetime.utcnow()
            generated = existing
        else:
            revisions = [{
                "version": 1,
                "content": content,
                "model": result.get("model_used"),
                "feedback": None,
                "timestamp": datetime.utcnow().isoformat(),
            }]
            generated = GeneratedCoverLetter(
                job_id=job_id,
                user_id=user_id,
                current_content=content,
                revisions=revisions,
            )
            db.add(generated)

        db.commit()
        db.refresh(generated)
        _set_generation_status(job_id, "completed", cover_letter_id=generated.id, version=len(generated.revisions), step="Done", percent=100)
    except Exception as e:
        logger.error(f"[generate-cover-letter-bg] Error for job {job_id}: {e}", exc_info=True)
        _set_generation_status(job_id, "error", error=str(e), step="Error", percent=0)
    finally:
        db.close()


@app.post("/api/jobs/{job_id}/generate-cover-letter")
async def generate_job_cover_letter(job_id: int, request: dict, background_tasks: BackgroundTasks, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """Generate a cover letter for a job. Runs in background; poll /generate-cover-letter/status."""
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == current_user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    job_description = job.job_description_parsed or job.job_description_raw

    # Pull the latest generated resume to use as content reference
    latest_resume = db.query(GeneratedResume).filter(GeneratedResume.job_id == job_id).order_by(GeneratedResume.updated_at.desc()).first()
    resume_content = latest_resume.current_content if latest_resume else None

    # Example cover letters from DB unless provided in request
    example_cover_letters = request.get("example_cover_letters", [])
    if not example_cover_letters:
        letters_db = db.query(BaseCoverLetter).filter(BaseCoverLetter.user_id == current_user.id).all()
        example_cover_letters = [{"name": l.name, "content": l.content} for l in letters_db]

    model_override = request.get("model")

    _set_generation_status(job_id, "processing", step="Queued", percent=0)
    background_tasks.add_task(
        _do_generate_cover_letter,
        job_id=job_id,
        user_id=current_user.id,
        job_description=job_description,
        resume_content=resume_content,
        example_cover_letters=example_cover_letters,
        target_role=job.position,
        company=job.company,
        position=job.position,
        model_override=model_override,
    )
    return {"job_id": job_id, "status": "processing"}


@app.get("/api/jobs/{job_id}/generate-cover-letter/status")
async def get_generate_cover_letter_status(job_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """Poll for background cover letter generation status."""
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == current_user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    status = _generation_status.get(job_id, {"status": "unknown"})
    if status.get("status") == "completed":
        latest = db.query(GeneratedCoverLetter).filter(GeneratedCoverLetter.job_id == job_id).order_by(GeneratedCoverLetter.updated_at.desc()).first()
        if latest:
            return {
                "status": "completed",
                "job_id": job_id,
                "cover_letter": latest.current_content,
                "cover_letter_id": latest.id,
                "version": len(latest.revisions) if latest.revisions else 1,
                "revisions": latest.revisions or [],
            }
    return status


@app.post("/api/jobs/{job_id}/revise-cover-letter")
async def revise_job_cover_letter(job_id: int, request: dict, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """Revise a generated cover letter based on user feedback. Appends a new revision."""
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == current_user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    feedback = request.get("feedback", "")
    if not feedback:
        raise HTTPException(status_code=400, detail="Feedback is required")

    existing = db.query(GeneratedCoverLetter).filter(GeneratedCoverLetter.job_id == job_id).first()
    if not existing or not existing.current_content:
        raise HTTPException(status_code=404, detail="No generated cover letter found. Generate one first.")

    # Pull the latest resume for fact reference
    latest_resume = db.query(GeneratedResume).filter(GeneratedResume.job_id == job_id).order_by(GeneratedResume.updated_at.desc()).first()
    resume_content = latest_resume.current_content if latest_resume else None

    model_override = request.get("model")
    result = await agent_service.revise_cover_letter(
        current_content=existing.current_content,
        feedback=feedback,
        job_description=job.job_description_parsed or job.job_description_raw,
        resume_content=resume_content,
        target_role=job.position,
        model_override=model_override,
    )

    if "error" in result and not result.get("content"):
        raise HTTPException(status_code=500, detail=result["error"])

    content = result.get("content", "")
    from sqlalchemy.orm.attributes import flag_modified
    revisions = list(existing.revisions or [])
    max_version = max((r.get("version", 0) for r in revisions if isinstance(r, dict)), default=0)
    next_version = max_version + 1
    revisions.append({
        "version": next_version,
        "content": content,
        "model": result.get("model_used"),
        "feedback": feedback,
        "timestamp": datetime.utcnow().isoformat(),
    })
    existing.current_content = content
    existing.revisions = revisions
    flag_modified(existing, "revisions")
    existing.updated_at = datetime.utcnow()
    db.commit()

    return {
        "job_id": job_id,
        "cover_letter": content,
        "cover_letter_id": existing.id,
        "version": next_version,
        "revisions": revisions,
    }


@app.delete("/api/jobs/{job_id}/cover-letter")
def delete_job_cover_letter(job_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """Delete the generated cover letter for a job."""
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == current_user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    existing = db.query(GeneratedCoverLetter).filter(GeneratedCoverLetter.job_id == job_id).first()
    if not existing:
        raise HTTPException(status_code=404, detail="No cover letter found")
    db.delete(existing)
    db.commit()
    return {"message": "Cover letter deleted"}


class CoverLetterDocxRequest(BaseModel):
    content: str
    version: Optional[int] = None  # if provided, export that specific revision


@app.post("/api/jobs/{job_id}/cover-letter/export-docx")
def export_cover_letter_docx(
    job_id: int,
    request: CoverLetterDocxRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Build a DOCX file from the cover letter plain text and return it as base64.

    Unlike resume DOCX export, cover letters have no template atoms — we
    render the plain text as a clean, single-column document with sensible
    defaults (Georgia 11pt, 1.15 line height, 1" margins).
    """
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == current_user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    content = (request.content or "").strip()
    if not content:
        raise HTTPException(status_code=400, detail="Cover letter content is empty")

    try:
        import io
        import base64
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_LINE_SPACING

        doc = Document()
        # Sensible defaults — Georgia 11pt is standard for cover letters.
        style = doc.styles["Normal"]
        style.font.name = "Georgia"
        style.font.size = Pt(11)
        # 1" margins all around.
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Split paragraphs on blank lines; preserve hard line breaks within a paragraph.
        for paragraph_text in content.split("\n\n"):
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(12)
            # 1.15 line spacing reads better for prose than the default 1.08.
            para.paragraph_format.line_spacing = 1.15
            lines = paragraph_text.split("\n")
            for i, line in enumerate(lines):
                if i > 0:
                    para.add_run().add_break()
                para.add_run(line)

        buf = io.BytesIO()
        doc.save(buf)
        out_bytes = buf.getvalue()

        version_num = request.version or 1
        filename = f"cover-letter-v{version_num}.docx"
        return {
            "docx_base64": base64.b64encode(out_bytes).decode("ascii"),
            "filename_suggestion": filename,
            "size_bytes": len(out_bytes),
        }
    except Exception as e:
        logger.error(f"cover letter export-docx failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to build DOCX: {e}")


# ---- Scoring agents: ATS + industry panel -----------------------------

async def _do_score(job_id: int, user_id: int, artifact_type: str, artifact_id: int, score_type: str, artifact_content: str, job_description: str, target_role: Optional[str]):
    """Background scoring (ATS or industry panel) for a resume or cover letter."""
    db = SessionLocal()
    try:
        _set_generation_status(job_id, "processing", step=f"Scoring ({score_type})", percent=10)
        if score_type == "ats":
            result = await agent_service.score_ats(artifact_type, artifact_content, job_description)
        else:
            result = await agent_service.score_industry_panel(artifact_type, artifact_content, job_description, target_role)

        score = ArtifactScore(
            job_id=job_id,
            user_id=user_id,
            artifact_type=artifact_type,
            artifact_id=artifact_id,
            score_type=score_type,
            scores=result.get("scores", {}),
            issues=result.get("issues", []),
            recommendations=result.get("recommendations", []),
            model_used=result.get("model_used"),
        )
        db.add(score)
        db.commit()
        db.refresh(score)
        _set_generation_status(job_id, "completed", score_id=score.id, score_type=score_type, step="Done", percent=100)
    except Exception as e:
        logger.error(f"[score-bg] Error {score_type} for job {job_id}: {e}", exc_info=True)
        _set_generation_status(job_id, "error", error=str(e), step="Error", percent=0)
    finally:
        db.close()


def _resolve_artifact_for_scoring(db: Session, job_id: int, artifact_type: str, artifact_id: Optional[int]):
    """Return (artifact_id, content) for the latest resume or cover letter on the job."""
    if artifact_type == "cover_letter":
        artifact = db.query(GeneratedCoverLetter).filter(GeneratedCoverLetter.job_id == job_id).order_by(GeneratedCoverLetter.updated_at.desc()).first()
    else:
        artifact = db.query(GeneratedResume).filter(GeneratedResume.job_id == job_id).order_by(GeneratedResume.updated_at.desc()).first()
    if not artifact:
        return None, None
    return (artifact_id or artifact.id), artifact.current_content


@app.post("/api/jobs/{job_id}/resumes/{resume_id}/score-ats")
async def score_resume_ats(job_id: int, resume_id: int, background_tasks: BackgroundTasks, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """Run ATS scoring on a resume."""
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == current_user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    resume = db.query(GeneratedResume).filter(GeneratedResume.id == resume_id, GeneratedResume.job_id == job_id).first()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    _set_generation_status(job_id, "processing", step="Scoring (ats)", percent=0)
    background_tasks.add_task(_do_score, job_id=job_id, user_id=current_user.id, artifact_type="resume", artifact_id=resume.id, score_type="ats", artifact_content=resume.current_content, job_description=job.job_description_parsed or job.job_description_raw, target_role=job.position)
    return {"job_id": job_id, "status": "processing", "score_type": "ats"}


@app.post("/api/jobs/{job_id}/resumes/{resume_id}/score-industry-panel")
async def score_resume_panel(job_id: int, resume_id: int, background_tasks: BackgroundTasks, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """Run industry panel scoring on a resume."""
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == current_user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    resume = db.query(GeneratedResume).filter(GeneratedResume.id == resume_id, GeneratedResume.job_id == job_id).first()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    _set_generation_status(job_id, "processing", step="Scoring (industry_panel)", percent=0)
    background_tasks.add_task(_do_score, job_id=job_id, user_id=current_user.id, artifact_type="resume", artifact_id=resume.id, score_type="industry_panel", artifact_content=resume.current_content, job_description=job.job_description_parsed or job.job_description_raw, target_role=job.position)
    return {"job_id": job_id, "status": "processing", "score_type": "industry_panel"}


@app.post("/api/jobs/{job_id}/cover-letters/{cl_id}/score-ats")
async def score_cover_letter_ats(job_id: int, cl_id: int, background_tasks: BackgroundTasks, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """Run ATS scoring on a cover letter."""
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == current_user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    cl = db.query(GeneratedCoverLetter).filter(GeneratedCoverLetter.id == cl_id, GeneratedCoverLetter.job_id == job_id).first()
    if not cl:
        raise HTTPException(status_code=404, detail="Cover letter not found")
    _set_generation_status(job_id, "processing", step="Scoring (ats)", percent=0)
    background_tasks.add_task(_do_score, job_id=job_id, user_id=current_user.id, artifact_type="cover_letter", artifact_id=cl.id, score_type="ats", artifact_content=cl.current_content, job_description=job.job_description_parsed or job.job_description_raw, target_role=job.position)
    return {"job_id": job_id, "status": "processing", "score_type": "ats"}


@app.post("/api/jobs/{job_id}/cover-letters/{cl_id}/score-industry-panel")
async def score_cover_letter_panel(job_id: int, cl_id: int, background_tasks: BackgroundTasks, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """Run industry panel scoring on a cover letter."""
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == current_user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    cl = db.query(GeneratedCoverLetter).filter(GeneratedCoverLetter.id == cl_id, GeneratedCoverLetter.job_id == job_id).first()
    if not cl:
        raise HTTPException(status_code=404, detail="Cover letter not found")
    _set_generation_status(job_id, "processing", step="Scoring (industry_panel)", percent=0)
    background_tasks.add_task(_do_score, job_id=job_id, user_id=current_user.id, artifact_type="cover_letter", artifact_id=cl.id, score_type="industry_panel", artifact_content=cl.current_content, job_description=job.job_description_parsed or job.job_description_raw, target_role=job.position)
    return {"job_id": job_id, "status": "processing", "score_type": "industry_panel"}


@app.get("/api/jobs/{job_id}/scores")
def get_job_scores(job_id: int, artifact_type: Optional[str] = None, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """Get the latest ATS + industry-panel scores for a job's resume and/or cover letter."""
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == current_user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")
    query = db.query(ArtifactScore).filter(ArtifactScore.job_id == job_id)
    if artifact_type:
        query = query.filter(ArtifactScore.artifact_type == artifact_type)
    scores = query.order_by(ArtifactScore.created_at.desc()).all()
    # Return latest per (artifact_type, artifact_id, score_type)
    seen = {}
    out = []
    for s in scores:
        key = (s.artifact_type, s.artifact_id, s.score_type)
        if key in seen:
            continue
        seen[key] = True
        out.append({
            "id": s.id,
            "artifact_type": s.artifact_type,
            "artifact_id": s.artifact_id,
            "score_type": s.score_type,
            "scores": s.scores,
            "issues": s.issues,
            "recommendations": s.recommendations,
            "model_used": s.model_used,
            "created_at": s.created_at.isoformat() if s.created_at else None,
        })
    return out


# ---- URL-only job creation (httpx first, fetcher sidecar fallback) ----

@app.post("/api/jobs/from-url")
async def create_job_from_url(request: dict, background_tasks: BackgroundTasks, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """Fetch a job posting from a URL, parse it, and create the job in one call.

    Tries direct httpx first; on 403/blocked or empty parse, falls back to the
    fetcher sidecar (FETCHER_URL env var, default http://fetcher:8080) which
    renders with a headless browser.
    """
    url = request.get("url")
    if not url:
        raise HTTPException(status_code=400, detail="URL is required")
    is_safe, reason = is_url_safe(url)
    if not is_safe:
        raise HTTPException(status_code=400, detail=f"URL is not safe to fetch: {reason}")

    headers = {
        "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.9",
    }

    html_content = None
    fetch_error = None

    # Attempt 1: direct httpx
    try:
        async with httpx.AsyncClient(follow_redirects=True, timeout=30.0) as client:
            response = await client.get(url, headers=headers)
            response.raise_for_status()
            html_content = response.text
    except httpx.HTTPStatusError as e:
        fetch_error = f"HTTP {e.response.status_code}"
    except httpx.RequestError as e:
        fetch_error = f"network: {str(e)}"
    except Exception as e:
        fetch_error = str(e)

    # Attempt 2: fetcher sidecar fallback (Playwright) on 403/blocked/empty
    fetcher_url = os.getenv("FETCHER_URL", "http://fetcher:8080")
    if (not html_content or fetch_error) and fetcher_url:
        try:
            async with httpx.AsyncClient(timeout=60.0) as client:
                r = await client.post(f"{fetcher_url}/fetch", json={"url": url}, headers={"Content-Type": "application/json"})
                r.raise_for_status()
                data = r.json()
                html_content = data.get("html")
                fetch_error = None
        except Exception as e:
            # If the sidecar isn't reachable, keep the original error
            if not fetch_error:
                fetch_error = f"fetcher sidecar: {str(e)}"

    if not html_content:
        raise HTTPException(status_code=502, detail=f"Failed to fetch URL: {fetch_error or 'empty response'}")

    parser = JobParser()
    try:
        job_data = await parser.parse_from_html(html_content, url)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to parse fetched page: {str(e)}")

    # Extract structured pay range + deadline
    pay_range = job_data.get("pay_range") or {}
    deadline_str = job_data.get("application_deadline")
    application_deadline = None
    if deadline_str:
        try:
            application_deadline = datetime.fromisoformat(deadline_str)
        except (ValueError, TypeError):
            application_deadline = None

    job = Job(
        public_job_id=_generate_unique_public_job_id(db),
        user_id=current_user.id,
        company=job_data.get("company", "Unknown"),
        position=job_data.get("position", "Unknown"),
        location=job_data.get("location"),
        salary=job_data.get("salary"),
        pay_range_min=pay_range.get("min"),
        pay_range_max=pay_range.get("max"),
        pay_currency=pay_range.get("currency", "USD"),
        pay_period=pay_range.get("period"),
        application_deadline=application_deadline,
        remote=job_data.get("remote"),
        job_url=url,
        job_description_raw=job_data.get("raw_text") or job_data.get("description"),
        job_description_parsed=job_data.get("description"),
        requirements=job_data.get("requirements", {}),
        responsibilities=job_data.get("responsibilities", []),
        keywords=job_data.get("keywords", []),
        required_credentials=job_data.get("credentials", []),
        source_type="url",
        source_url=url,
    )
    db.add(job)
    db.commit()
    db.refresh(job)

    application = JobApplication(
        job_id=job.id,
        user_id=current_user.id,
        stage="saved",
        history=[{"date": datetime.utcnow().isoformat(), "action": "job_added", "notes": "Job added from URL"}],
    )
    db.add(application)
    db.commit()

    return format_job_response(job, application, db)


# Agent endpoints
@app.post("/api/agents/ats-analysis")
async def ats_analysis(resume_text: str, job_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """Run ATS Expert Agent analysis on a resume against a job"""
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == current_user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    job_description = job.job_description_parsed or job.job_description_raw or ""
    analysis = agent_service.get_ats_expert_analysis(resume_text, job_description)

    return {"job_id": job_id, "analysis": analysis}


@app.post("/api/agents/technical-fit")
async def technical_fit(resume_text: str, job_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """Run Technical Hiring Manager Agent analysis"""
    job = db.query(Job).filter(Job.id == job_id, Job.user_id == current_user.id).first()
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    analysis = agent_service.get_technical_fit_analysis(
        resume_text, job.requirements or {}
    )

    return {"job_id": job_id, "analysis": analysis}


@app.post("/api/agents/generate-resume")
async def generate_resume(
    user_profile: dict, job_id: Optional[int] = None, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)
):
    """Generate a resume using Resume Generator Agent"""
    job_description = None
    if job_id:
        job = db.query(Job).filter(Job.id == job_id, Job.user_id == current_user.id).first()
        if job:
            job_description = job.job_description_parsed or job.job_description_raw

    resume = await agent_service.generate_resume(
        user_profile,
        job_description=job_description,
        target_role=user_profile.get("target_role"),
    )

    return {"job_id": job_id, "resume": resume}


# Merged MCP routes (previously mcp_server.py)


@app.post("/api/fetch-job")
async def fetch_job(request: dict, current_user: User = Depends(get_current_user)):
    """
    Fetch job details from a URL.
    Supports LinkedIn, Indeed, and generic job boards.
    Requires authentication to prevent abuse as an open proxy.
    """
    url = request.get("url")
    if not url:
        raise HTTPException(status_code=400, detail="URL is required")

    # SSRF protection: validate URL before fetching
    is_safe, reason = is_url_safe(url)
    if not is_safe:
        raise HTTPException(status_code=400, detail=f"URL is not safe to fetch: {reason}")

    headers = {
        "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.9",
        "Accept-Encoding": "gzip, deflate, br",
        "DNT": "1",
        "Connection": "keep-alive",
        "Upgrade-Insecure-Requests": "1",
        "Sec-Fetch-Dest": "document",
        "Sec-Fetch-Mode": "navigate",
        "Sec-Fetch-Site": "none",
        "Sec-Fetch-User": "?1",
        "Cache-Control": "max-age=0",
    }

    try:
        async with httpx.AsyncClient(follow_redirects=True, timeout=30.0) as client:
            response = await client.get(url, headers=headers)
            response.raise_for_status()
            html_content = response.text

        parser = JobParser()
        job_data = await parser.parse_from_html(html_content, url)
        return job_data

    except httpx.HTTPStatusError as e:
        if e.response.status_code == 403:
            raise HTTPException(
                status_code=403,
                detail=f"Access denied (403) when fetching URL. The website may block automated requests: {url}",
            )
        elif e.response.status_code == 404:
            raise HTTPException(
                status_code=404,
                detail=f"Job posting not found (404): {url}",
            )
        else:
            raise HTTPException(
                status_code=502,
                detail=f"Failed to fetch URL (HTTP {e.response.status_code}): {str(e)}",
            )
    except httpx.RequestError as e:
        raise HTTPException(
            status_code=502,
            detail=f"Network error when fetching URL: {str(e)}",
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing job: {str(e)}")


# Static file serving and SPA catch-all
STATIC_DIR = str(Path(__file__).parent.parent / "static")
app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")


# Brand assets served at root (favicon + logo referenced from index.html <head>/<header>)
@app.get("/joblign-favicon.png")
async def brand_favicon_png():
    return FileResponse(Path(STATIC_DIR) / "joblign-favicon.png", media_type="image/png")


@app.get("/joblign-favicon.webp")
async def brand_favicon_webp():
    return FileResponse(Path(STATIC_DIR) / "joblign-favicon.webp", media_type="image/webp")


@app.get("/joblign-logo.png")
async def brand_logo_png():
    return FileResponse(Path(STATIC_DIR) / "joblign-logo.png", media_type="image/png")


@app.get("/joblign-logo.webp")
async def brand_logo_webp():
    return FileResponse(Path(STATIC_DIR) / "joblign-logo.webp", media_type="image/webp")


@app.get("/")
async def serve_index():
    """Serve the main frontend page"""
    return FileResponse(Path(STATIC_DIR) / "index.html")


@app.get("/{full_path:path}")
async def serve_frontend(full_path: str):
    """Catch-all for SPA routing — return index.html for any non-API path"""
    if full_path.startswith("api/"):
        raise HTTPException(status_code=404, detail="Not found")
    return FileResponse(Path(STATIC_DIR) / "index.html")


if __name__ == "__main__":
    import uvicorn

    port = int(os.getenv("PORT", "8765"))
    uvicorn.run(app, host="0.0.0.0", port=port)
