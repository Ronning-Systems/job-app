"""
Template parser and DOCX composer for template-driven resume generation.

Pipeline:
  1. parse_template(docx_bytes) -> atoms dict
       Reads a user's template DOCX, classifies each paragraph by structural
       role (TITLE, SECTION_HEADER, ROLE_TITLE, ROLE_LINE, BULLET, BODY_PARA)
       based on the template's own numbering definitions and style usage,
       captures every relevant style attribute (font, size, color, bold,
       italic, spacing, indent, bullet numId/ilvl) into an atom, and
       persists the atoms alongside the raw DOCX so the composer can later
       reproduce the exact look.

  2. compose_docx(template_docx_bytes, atoms, structured_content) -> docx_bytes
       Loads the original template DOCX (so its styles.xml, numbering.xml,
       theme1.xml etc. survive verbatim), wipes its body content, then
       rebuilds paragraphs from structured_content applying each paragraph's
       atom style. Inline runs (bold/italic toggles, segments separated by
       pipes, etc.) are honored.

The point: the LLM never has to invent font sizes, hex colors, or spacing.
It only decides what text goes where. Visual fidelity comes from the captured
template atoms.
"""

from __future__ import annotations

import base64
import copy
import io
import logging
from dataclasses import dataclass, field, asdict
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.oxml.ns import qn
from lxml import etree

logger = logging.getLogger(__name__)

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W = "{" + W_NS + "}"


def _w(tag: str) -> str:
    return f"{W}{tag}"


# ---------- Atom dataclass --------------------------------------------------


@dataclass
class Atom:
    """A reusable style atom captured from the template.

    One atom per distinct *visual* role in the resume (title, section header,
    bullet, etc.). All fields are denormalized so the composer doesn't need
    to chase down style inheritance at render time.
    """

    id: str
    docx_style: str              # python-docx style name (e.g., "Body Text", "Heading 2")
    style_id: str                # raw styleId (e.g., "BodyText", "Heading2")
    font_ascii: Optional[str] = None
    font_h_ansi: Optional[str] = None
    size_half_points: Optional[int] = None
    bold: Optional[bool] = None
    italic: Optional[bool] = None
    underline: Optional[bool] = None
    color_hex: Optional[str] = None     # 6-char hex without '#', or None for auto
    spacing_before_twips: Optional[int] = None
    spacing_after_twips: Optional[int] = None
    line_spacing_240ths: Optional[int] = None  # 240 = single, 276 = 1.15, 360 = 1.5
    line_rule: Optional[str] = None     # "auto", "exact", "atLeast"
    alignment: Optional[str] = None    # "start", "center", "end", "both"
    indent_left_twips: Optional[int] = None
    indent_right_twips: Optional[int] = None
    indent_first_line_twips: Optional[int] = None
    indent_hanging_twips: Optional[int] = None
    num_id: Optional[int] = None
    num_ilvl: Optional[int] = None
    description: str = ""
    sample_text: str = ""

    def to_dict(self) -> Dict[str, Any]:
        return {k: v for k, v in asdict(self).items() if v is not None or k in {"id", "docx_style", "style_id"}}


# ---------- Template parser --------------------------------------------------


# Canonical atom roles. Each role can have multiple "variants" — when the
# template uses more than one visual style for the same conceptual role
# (e.g., summary uses decimal-numbered "1) 2) 3)" items while the experience
# bullets use a different bullet character), we capture each variant as a
# separate atom. The LLM picks which variant to use by atom_id.
CANONICAL_ATOMS = [
    "title",            # Big name + tagline at top (single paragraph, can contain separator runs)
    "section_header",   # "Core Competencies", "Professional Experience", etc.
    "role_title",       # "Chief Technology Officer"
    "role_line",        # "Manna Health AI | Toronto, Ontario, Canada | March 2026 - Present"
    "bullet",           # A single accomplishment bullet
    "body_para",        # Summary paragraph(s) and other plain prose
]


def _get_style_def(doc: Document, style_id: str) -> Optional[etree._Element]:
    """Return the <w:style> element with the given styleId, or None."""
    styles_el = doc.styles.element
    for st in styles_el.findall(_w("style")):
        if st.get(_w("styleId")) == style_id:
            return st
    return None


def _resolve_run_font(rpr: etree._Element) -> Tuple[Optional[str], Optional[str]]:
    """Return (ascii_font, h_ansi_font) from a <w:rPr>."""
    rfonts = rpr.find(_w("rFonts"))
    if rfonts is None:
        return None, None
    return rfonts.get(_w("ascii")), rfonts.get(_w("hAnsi"))


def _resolve_size_hp(rpr: etree._Element) -> Optional[int]:
    """Resolve font size (in half-points) from a <w:rPr>, walking style chain."""
    sz = rpr.find(_w("sz"))
    if sz is not None:
        return int(sz.get(_w("val")))
    return None


def _resolve_color(rpr: etree._Element) -> Optional[str]:
    color = rpr.find(_w("color"))
    if color is None:
        return None
    val = color.get(_w("val"))
    if not val or val == "auto":
        return None
    return val.upper()


def _resolve_bool(rpr: etree._Element, tag: str) -> Optional[bool]:
    el = rpr.find(_w(tag))
    if el is None:
        return None
    val = el.get(_w("val"))
    if val is None:
        return True  # bare <w:b/> means true
    return val.lower() not in ("0", "false")


def _resolve_spacing(ppr: etree._Element) -> Tuple[Optional[int], Optional[int], Optional[int], Optional[str]]:
    """Return (before_twips, after_twips, line_240ths, line_rule)."""
    if ppr is None:
        return None, None, None, None
    sp = ppr.find(_w("spacing"))
    if sp is None:
        return None, None, None, None
    before = sp.get(_w("before"))
    after = sp.get(_w("after"))
    line = sp.get(_w("line"))
    rule = sp.get(_w("lineRule"))
    return (
        int(before) if before else None,
        int(after) if after else None,
        int(line) if line else None,
        rule,
    )


def _resolve_indent(ppr: etree._Element) -> Tuple[Optional[int], Optional[int], Optional[int], Optional[int]]:
    """Return (left, right, first_line, hanging) twips."""
    if ppr is None:
        return None, None, None, None
    ind = ppr.find(_w("ind"))
    if ind is None:
        return None, None, None, None
    # 'start' is the modern (theme-aware) left; 'left' is the legacy equivalent
    left = ind.get(_w("start")) or ind.get(_w("left"))
    right = ind.get(_w("end")) or ind.get(_w("right"))
    first = ind.get(_w("firstLine"))
    hanging = ind.get(_w("hanging"))
    return (
        int(left) if left else None,
        int(right) if right else None,
        int(first) if first else None,
        int(hanging) if hanging else None,
    )


def _resolve_alignment(ppr: etree._Element) -> Optional[str]:
    if ppr is None:
        return None
    jc = ppr.find(_w("jc"))
    if jc is None:
        return None
    return jc.get(_w("val"))


def _resolve_numbering(ppr: etree._Element) -> Tuple[Optional[int], Optional[int]]:
    """Return (num_id, ilvl)."""
    if ppr is None:
        return None, None
    num_pr = ppr.find(_w("numPr"))
    if num_pr is None:
        return None, None
    ilvl_el = num_pr.find(_w("ilvl"))
    num_id_el = num_pr.find(_w("numId"))
    return (
        int(num_id_el.get(_w("val"))) if num_id_el is not None else None,
        int(ilvl_el.get(_w("val"))) if ilvl_el is not None else None,
    )


def _get_para_style_id(p: Any) -> Optional[str]:
    """Return the styleId (e.g., 'Heading2') for a paragraph, or None."""
    ppr = p._p.find(_w("pPr"))
    if ppr is None:
        return None
    pstyle = ppr.find(_w("pStyle"))
    if pstyle is None:
        return None
    return pstyle.get(_w("val"))


def _find_first_rpr_with_size(doc: Document, style_id: str) -> Optional[int]:
    """Walk basedOn chain looking for a font size in half-points."""
    seen = set()
    sid = style_id
    while sid and sid not in seen:
        seen.add(sid)
        st = _get_style_def(doc, sid)
        if st is None:
            return None
        rpr = st.find(_w("rPr"))
        if rpr is not None:
            sz = _resolve_size_hp(rpr)
            if sz is not None:
                return sz
        # walk basedOn
        based = st.find(_w("basedOn"))
        sid = based.get(_w("val")) if based is not None else None
    return None


def _find_first_rpr_with_font(doc: Document, style_id: str) -> Tuple[Optional[str], Optional[str]]:
    seen = set()
    sid = style_id
    while sid and sid not in seen:
        seen.add(sid)
        st = _get_style_def(doc, sid)
        if st is None:
            return None, None
        rpr = st.find(_w("rPr"))
        if rpr is not None:
            a, h = _resolve_run_font(rpr)
            if a or h:
                return a, h
        based = st.find(_w("basedOn"))
        sid = based.get(_w("val")) if based is not None else None
    return None, None


def _find_first_rpr_with_color(doc: Document, style_id: str) -> Optional[str]:
    seen = set()
    sid = style_id
    while sid and sid not in seen:
        seen.add(sid)
        st = _get_style_def(doc, sid)
        if st is None:
            return None
        rpr = st.find(_w("rPr"))
        if rpr is not None:
            c = _resolve_color(rpr)
            if c:
                return c
        based = st.find(_w("basedOn"))
        sid = based.get(_w("val")) if based is not None else None
    return None


def _gather_runs_text_and_format(p: Any) -> Tuple[List[Tuple[str, Dict[str, Any]]], bool]:
    """Walk a paragraph's runs and return [(text, fmt_dict)] plus a 'has_any_formatting' flag.

    fmt_dict can contain keys: bold, italic, underline. Text is the run's
    visible string. Empty runs are skipped. Hyperlink fields are flattened
    to their inner text with formatting preserved from the first run inside.
    """
    out: List[Tuple[str, Dict[str, Any]]] = []
    has_fmt = False
    pPr = p._p.find(_w("pPr"))
    # paragraph-level rPr (the "end of paragraph" run formatting) — applied
    # implicitly to the trailing newline; rarely matters but capture it.
    for r in p._p.findall(_w("r")):
        rpr = r.find(_w("rPr"))
        text = "".join(t.text or "" for t in r.findall(_w("t")))
        if not text:
            continue
        fmt: Dict[str, Any] = {}
        if rpr is not None:
            b = _resolve_bool(rpr, "b")
            i = _resolve_bool(rpr, "i")
            u = _resolve_bool(rpr, "u")
            if b:
                fmt["bold"] = True
            if i:
                fmt["italic"] = True
            if u:
                fmt["underline"] = True
            if fmt:
                has_fmt = True
        out.append((text, fmt))
    # Hyperlinks: <w:hyperlink><w:r>...</w:r></w:hyperlink>
    for h in p._p.findall(_w("hyperlink")):
        # Hyperlinks inherit formatting from style + their own rStyle; we keep
        # any explicit b/i/u on inner runs.
        for r in h.findall(_w("r")):
            rpr = r.find(_w("rPr"))
            text = "".join(t.text or "" for t in r.findall(_w("t")))
            if not text:
                continue
            fmt: Dict[str, Any] = {}
            if rpr is not None:
                b = _resolve_bool(rpr, "b")
                i = _resolve_bool(rpr, "i")
                u = _resolve_bool(rpr, "u")
                if b:
                    fmt["bold"] = True
                if i:
                    fmt["italic"] = True
                if u:
                    fmt["underline"] = True
                if fmt:
                    has_fmt = True
            out.append((text, fmt))
    return out, has_fmt


def _classify_template_paragraphs(doc: Document) -> Tuple[Dict[int, str], Dict[str, List[Dict[str, Any]]]]:
    """Heuristically classify each paragraph in the template into one of the six
    canonical atom roles, AND collect distinct bullet variants.

    Returns (classification, variants) where:
      classification: {paragraph_index: atom_id} — atom_id may include a
        variant suffix like "bullet.section_list" or "bullet.role_list" when
        a role has multiple visual variants.
      variants: {atom_id: [{num_id, num_ilvl, source_para_idx, context, ...}]}
        — distinct style signatures observed for multi-variant atoms.
    """
    classification: Dict[int, str] = {}
    paras = doc.paragraphs

    # First pass: gather each paragraph's facts
    facts = []
    for i, p in enumerate(paras):
        sid = _get_para_style_id(p)
        ppr = p._p.find(_w("pPr"))
        num_id, num_ilvl = _resolve_numbering(ppr)
        bold_para = False
        if ppr is not None:
            rpr_pp = ppr.find(_w("rPr"))
            if rpr_pp is not None:
                bold_para = bool(_resolve_bool(rpr_pp, "b"))
        runs, has_inline_fmt = _gather_runs_text_and_format(p)
        facts.append({
            "idx": i,
            "style_id": sid,
            "style_name": p.style.name if p.style else None,
            "num_id": num_id,
            "num_ilvl": num_ilvl,
            "bold_para": bold_para,
            "text": p.text or "",
            "has_inline_fmt": has_inline_fmt,
            "runs": runs,
        })

    # Find the title paragraph: first H2 whose text contains "|" (separated
    # sub-title) OR the very first H2 if there is no other H2 before it.
    title_idx = None
    for f in facts:
        if f["style_id"] in ("Heading2",):
            if "|" in f["text"] or title_idx is None:
                title_idx = f["idx"]
                break
    if title_idx is None and facts and facts[0]["style_id"] in ("Heading2",):
        title_idx = 0

    for f in facts:
        idx = f["idx"]
        sid = f["style_id"]
        if idx == title_idx:
            classification[idx] = "title"
            continue
        if sid == "Heading2":
            classification[idx] = "section_header"
            continue
        if sid == "Heading3":
            classification[idx] = "role_title"
            continue
        if sid in ("BodyText", None) and f["num_id"] is not None:
            classification[idx] = "bullet"
            continue
        # Body text right after a role_title with "|" in it = role_line
        if sid in ("BodyText", None) and idx > 0:
            prev = facts[idx - 1]
            if prev["style_id"] == "Heading3" and "|" in f["text"]:
                classification[idx] = "role_line"
                continue
        classification[idx] = "body_para"

    # Re-walk: any "bullet" that is the first body-with-list after a
    # role_title AND has "|" -> demote to role_line (some templates use
    # list-formatting for the role_line too).
    for f in facts:
        idx = f["idx"]
        if classification.get(idx) == "bullet" and "|" in f["text"] and idx > 0:
            prev = facts[idx - 1]
            if prev["style_id"] == "Heading3":
                classification[idx] = "role_line"

    # ---- Multi-variant bullet detection --------------------------------
    # Walk through and group bullets by their (num_id, num_ilvl, after-style)
    # signature, where "after-style" means "what type of parent came right
    # before this bullet run": section_header (immediately preceded by an
    # H2), role_line (immediately preceded by a role_line), body_para (in
    # summary area), or other.
    bullet_groups: Dict[Tuple, List[int]] = {}
    for i, f in enumerate(facts):
        if classification.get(f["idx"]) != "bullet":
            continue
        # Find the most recent non-bullet ancestor
        context = "other"
        for j in range(i - 1, -1, -1):
            cl = classification.get(facts[j]["idx"])
            if cl == "section_header":
                context = "section_list"
                break
            if cl == "role_line":
                context = "role_list"
                break
            if cl == "role_title":
                context = "role_list"
                break
            if cl == "body_para":
                context = "summary_list"
                break
            if cl in ("title",):
                context = "summary_list"
                break
        key = (f["num_id"], f["num_ilvl"], context)
        bullet_groups.setdefault(key, []).append(f["idx"])

    variants: Dict[str, List[Dict[str, Any]]] = {"bullet": []}
    variant_id_map: Dict[Tuple, str] = {}
    # Sort by priority: summary_list, section_list, role_list, other
    priority_order = {"summary_list": 0, "section_list": 1, "role_list": 2, "other": 3}
    sorted_keys = sorted(bullet_groups.keys(), key=lambda k: priority_order.get(k[2], 99))

    suffix_lookup = {"summary_list": "summary", "section_list": "section", "role_list": "role", "other": "misc"}
    for key in sorted_keys:
        nid, ilvl, ctx = key
        suffix = suffix_lookup.get(ctx, "misc")
        atom_id_full = f"bullet.{suffix}"
        variant_id_map[key] = atom_id_full
        # Pick the FIRST paragraph in the group as the representative
        para_idx = bullet_groups[key][0]
        variants["bullet"].append({
            "atom_id": atom_id_full,
            "num_id": nid,
            "num_ilvl": ilvl,
            "source_para_idx": para_idx,
            "context": ctx,
            "count_in_template": len(bullet_groups[key]),
        })
        # Re-tag every paragraph in this group
        for pi in bullet_groups[key]:
            classification[pi] = atom_id_full

    return classification, variants


def _atom_from_paragraph(
    doc: Document,
    p: Any,
    atom_id: str,
    style_id: Optional[str],
) -> Atom:
    """Build an Atom from a paragraph and its style chain."""
    ppr = p._p.find(_w("pPr"))
    # Defaults come from the paragraph itself; fall back to style chain.
    before, after, line, line_rule = _resolve_spacing(ppr)
    left, right, first, hanging = _resolve_indent(ppr)
    align = _resolve_alignment(ppr)
    num_id, num_ilvl = _resolve_numbering(ppr)

    # Run-level: capture the first run's formatting as a representative
    # sample (paragraphs in the template usually have uniform run formatting
    # for atoms like title/role_title/section_header; bullets/role_lines have
    # inline toggles which the structured JSON will preserve separately).
    first_rpr = None
    for r in p._p.findall(_w("r")):
        rpr = r.find(_w("rPr"))
        if rpr is not None:
            first_rpr = rpr
            break

    font_a, font_h = (None, None)
    size_hp = None
    color = None
    bold = None
    italic = None
    underline = None
    if first_rpr is not None:
        font_a, font_h = _resolve_run_font(first_rpr)
        size_hp = _resolve_size_hp(first_rpr)
        color = _resolve_color(first_rpr)
        bold = _resolve_bool(first_rpr, "b")
        italic = _resolve_bool(first_rpr, "i")
        underline = _resolve_bool(first_rpr, "u")

    # Walk style chain for missing fields
    if style_id:
        if size_hp is None:
            size_hp = _find_first_rpr_with_size(doc, style_id)
        if not font_a and not font_h:
            font_a, font_h = _find_first_rpr_with_font(doc, style_id)
        if color is None:
            color = _find_first_rpr_with_color(doc, style_id)

    # Override with explicit paragraph numbering's bullet format if present
    # (we don't propagate the numbering style into the atom; the composer
    # looks up the abstractNum to determine the bullet character).
    bullet_text = ""
    if num_id is not None:
        try:
            numbering = doc.part.numbering_part.element
            num_el = None
            for n in numbering.findall(_w("num")):
                if int(n.get(_w("numId"))) == num_id:
                    num_el = n
                    break
            if num_el is not None:
                abs_id = int(num_el.find(_w("abstractNumId")).get(_w("val")))
                for abs_num in numbering.findall(_w("abstractNum")):
                    if int(abs_num.get(_w("abstractNumId"))) == abs_id:
                        ilvl_target = num_ilvl if num_ilvl is not None else 0
                        for lvl in abs_num.findall(_w("lvl")):
                            if int(lvl.get(_w("ilvl"))) == ilvl_target:
                                lvl_text_el = lvl.find(_w("lvlText"))
                                lvl_fmt_el = lvl.find(_w("numFmt"))
                                if lvl_text_el is not None:
                                    bullet_text = lvl_text_el.get(_w("val")) or ""
                                if lvl_fmt_el is not None and lvl_fmt_el.get(_w("val")) == "none":
                                    bullet_text = ""
                                break
                        break
        except Exception as e:
            logger.debug(f"Could not resolve bullet text for numId={num_id}: {e}")

    return Atom(
        id=atom_id,
        docx_style=p.style.name if p.style else "",
        style_id=style_id or "",
        font_ascii=font_a,
        font_h_ansi=font_h,
        size_half_points=size_hp,
        bold=bold,
        italic=italic,
        underline=underline,
        color_hex=color,
        spacing_before_twips=before,
        spacing_after_twips=after,
        line_spacing_240ths=line,
        line_rule=line_rule,
        alignment=align,
        indent_left_twips=left,
        indent_right_twips=right,
        indent_first_line_twips=first,
        indent_hanging_twips=hanging,
        num_id=num_id,
        num_ilvl=num_ilvl,
        description=ATOM_DESCRIPTIONS.get(atom_id, ""),
        sample_text=(p.text or "")[:120],
    )


ATOM_DESCRIPTIONS = {
    "title": "Top-of-resume name + tagline in a single paragraph (segments typically separated by ' | ')",
    "section_header": "Section heading (e.g., 'Core Competencies', 'Professional Experience')",
    "role_title": "Job title within a role entry (e.g., 'Chief Technology Officer')",
    "role_line": "Company, location, dates line directly under a role title (segments typically separated by ' | ')",
    "bullet": "One accomplishment bullet under a role",
    "body_para": "Plain prose paragraph (summary text, etc.)",
}


def parse_template(docx_bytes: bytes) -> Dict[str, Any]:
    """Parse a user's template DOCX into the atoms representation.

    Returns a dict with keys:
      - atoms: list of {id, ...} dicts (one per atom role; multi-variant
               atoms like 'bullet' expand to bullet.summary / bullet.section
               / bullet.role based on observed template usage)
      - docx_base64: the raw DOCX bytes (base64) so the composer can rebuild
                     from the original styles.xml/numbering.xml
      - page_setup: dict with page size + margins
      - canonical_atoms_present: which atom ids were observed
      - warnings: any issues encountered
    """
    doc = Document(io.BytesIO(docx_bytes))
    classification, variants = _classify_template_paragraphs(doc)

    # Find one representative paragraph per atom_id; prefer the first one
    # we encounter.
    by_atom: Dict[str, Any] = {}
    for i, p in enumerate(doc.paragraphs):
        atom_id = classification.get(i)
        if atom_id and atom_id not in by_atom:
            by_atom[atom_id] = (i, p)

    atoms_out: List[Dict[str, Any]] = []
    warnings: List[str] = []
    seen_ids: set = set()

    # 1) Emit one atom per role. For bullet, multiple variants may exist —
    #    we already turned them into separate atom_ids in classification.
    for atom_id in CANONICAL_ATOMS:
        if atom_id in by_atom and atom_id not in seen_ids:
            _, p = by_atom[atom_id]
            style_id = _get_para_style_id(p)
            atom = _atom_from_paragraph(doc, p, atom_id, style_id)
            atoms_out.append(atom.to_dict())
            seen_ids.add(atom_id)

    # 2) Emit extra bullet variants (skipping the first one, which we
    #    already added above under id="bullet"). The first variant's id
    #    might have been emitted as plain "bullet"; the rest are "bullet.X".
    for variant in variants.get("bullet", []):
        if variant["atom_id"] in seen_ids:
            continue
        para_idx = variant["source_para_idx"]
        p = doc.paragraphs[para_idx]
        style_id = _get_para_style_id(p)
        atom = _atom_from_paragraph(doc, p, variant["atom_id"], style_id)
        atoms_out.append(atom.to_dict())
        seen_ids.add(variant["atom_id"])

    # 3) Synthesize any canonical atom that wasn't seen at all. Skip 'bullet'
    #    because if we have *any* bullet variant we already cover the role —
    #    a generic 'bullet' fallback would just be redundant.
    for atom_id in CANONICAL_ATOMS:
        if atom_id not in seen_ids:
            if atom_id == "bullet" and any(a.startswith("bullet.") for a in seen_ids):
                # Skip — we already have at least one specific bullet variant.
                continue
            fallback_atom = _synthesize_missing_atom(doc, atom_id)
            if fallback_atom:
                atoms_out.append(fallback_atom.to_dict())
                seen_ids.add(atom_id)
                warnings.append(f"Atom '{atom_id}' not present in template; synthesized from nearest style")

    # Page setup
    sec = doc.sections[0]
    page_setup = {
        "page_width_emu": int(sec.page_width) if sec.page_width else None,
        "page_height_emu": int(sec.page_height) if sec.page_height else None,
        "top_margin_emu": int(sec.top_margin) if sec.top_margin else None,
        "right_margin_emu": int(sec.right_margin) if sec.right_margin else None,
        "bottom_margin_emu": int(sec.bottom_margin) if sec.bottom_margin else None,
        "left_margin_emu": int(sec.left_margin) if sec.left_margin else None,
    }

    return {
        "atoms": atoms_out,
        "docx_base64": base64.b64encode(docx_bytes).decode("ascii"),
        "page_setup": page_setup,
        "canonical_atoms_present": sorted([a["id"] for a in atoms_out]),
        "warnings": warnings,
    }


def _synthesize_missing_atom(doc: Document, atom_id: str) -> Optional[Atom]:
    """Create an atom by cloning the closest existing style when the template
    doesn't have an example for a given atom role.
    """
    fallback_map = {
        "title": "Heading2",
        "section_header": "Heading2",
        "role_title": "Heading3",
        "role_line": "BodyText",
        "bullet": "BodyText",
        "body_para": "BodyText",
    }
    target_sid = fallback_map.get(atom_id)
    if not target_sid:
        return None
    # Find a paragraph using this style
    for p in doc.paragraphs:
        sid = _get_para_style_id(p)
        if sid == target_sid:
            return _atom_from_paragraph(doc, p, atom_id, target_sid)
    return None


# ---------- DOCX composer --------------------------------------------------


def compose_docx(
    template_docx_bytes: bytes,
    atoms: List[Dict[str, Any]],
    structured_content: Dict[str, Any],
) -> bytes:
    """Compose a DOCX from the template's styles + a structured content tree.

    Args:
        template_docx_bytes: raw bytes of the original template DOCX. We
            open this so styles.xml, numbering.xml, theme1.xml, fontTable.xml
            all survive verbatim — the only thing we throw away is the body.
        atoms: list of atom dicts (as produced by parse_template)
        structured_content: {
            "atoms": [
                {"atom_id": "title", "segments": [{"text": "Foo"}, {"text": " | "}, ...]},
                {"atom_id": "section_header", "text": "Core Competencies"},
                {"atom_id": "body_para", "segments": [{"text": "..."}]},
                {"atom_id": "role_title", "text": "Chief Technology Officer"},
                {"atom_id": "role_line", "segments": [
                    {"text": "Manna Health AI", "bold": true},
                    {"text": " | Toronto, Ontario, Canada | March 2026 - Present"}
                ]},
                {"atom_id": "bullet", "segments": [{"text": "..."}]},
                ...
            ]
        }

    Returns:
        Bytes of the composed DOCX file.
    """
    atom_by_id = {a["id"]: a for a in atoms}

    # Open the template, then strip its body
    doc = Document(io.BytesIO(template_docx_bytes))
    body = doc.element.body
    # Preserve final sectPr; remove every other direct child of body
    sect_pr = body.find(_w("sectPr"))
    for child in list(body):
        if child.tag != _w("sectPr"):
            body.remove(child)

    # Append new paragraphs
    for entry in structured_content.get("atoms", []):
        atom_id = entry.get("atom_id", "")
        atom = atom_by_id.get(atom_id)
        if atom is None:
            # Unknown atom_id — render as plain body_para to avoid breaking
            logger.warning(f"compose_docx: unknown atom_id '{atom_id}'; rendering as body_para")
            atom = atom_by_id.get("body_para", {"docx_style": "Normal", "style_id": "Normal"})

        para_el = _build_paragraph(atom, entry)
        if sect_pr is not None:
            sect_pr.addprevious(para_el)
        else:
            body.append(para_el)

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


def _build_paragraph(atom: Dict[str, Any], entry: Dict[str, Any]) -> etree._Element:
    """Build a <w:p> element applying atom's style + entry's text."""
    p = etree.Element(_w("p"))
    ppr = etree.SubElement(p, _w("pPr"))

    # pStyle — preserve the style the template used
    style_id = atom.get("style_id")
    if style_id:
        pstyle = etree.SubElement(ppr, _w("pStyle"))
        pstyle.set(_w("val"), style_id)

    # Numbering — only emit <w:numPr> when the atom has a real, positive
    # num_id. The template often uses numId=0 as a "no list" sentinel
    # (e.g., for the title row) and we must not propagate that into the
    # composed document or it would try to reference an undefined list.
    num_id = atom.get("num_id")
    num_ilvl = atom.get("num_ilvl")
    if num_id is not None and int(num_id) > 0:
        num_pr = etree.SubElement(ppr, _w("numPr"))
        ilvl_el = etree.SubElement(num_pr, _w("ilvl"))
        ilvl_el.set(_w("val"), str(num_ilvl if num_ilvl is not None else 0))
        num_id_el = etree.SubElement(num_pr, _w("numId"))
        num_id_el.set(_w("val"), str(int(num_id)))

    # Spacing — only set explicit values if the atom defines them; otherwise
    # let the style cascade handle it.
    has_explicit_spacing = (
        atom.get("spacing_before_twips") is not None
        or atom.get("spacing_after_twips") is not None
        or atom.get("line_spacing_240ths") is not None
    )
    if has_explicit_spacing:
        sp = etree.SubElement(ppr, _w("spacing"))
        if atom.get("spacing_before_twips") is not None:
            sp.set(_w("before"), str(atom["spacing_before_twips"]))
        if atom.get("spacing_after_twips") is not None:
            sp.set(_w("after"), str(atom["spacing_after_twips"]))
        if atom.get("line_spacing_240ths") is not None:
            sp.set(_w("line"), str(atom["line_spacing_240ths"]))
            rule = atom.get("line_rule") or "auto"
            sp.set(_w("lineRule"), rule)

    # Indent
    if any(atom.get(k) is not None for k in ("indent_left_twips", "indent_right_twips", "indent_first_line_twips", "indent_hanging_twips")):
        ind = etree.SubElement(ppr, _w("ind"))
        if atom.get("indent_left_twips") is not None:
            # Use 'start' for theme-aware left
            ind.set(_w("start"), str(atom["indent_left_twips"]))
            ind.set(_w("left"), str(atom["indent_left_twips"]))
        if atom.get("indent_right_twips") is not None:
            ind.set(_w("end"), str(atom["indent_right_twips"]))
            ind.set(_w("right"), str(atom["indent_right_twips"]))
        if atom.get("indent_first_line_twips") is not None:
            ind.set(_w("firstLine"), str(atom["indent_first_line_twips"]))
        if atom.get("indent_hanging_twips") is not None:
            ind.set(_w("hanging"), str(atom["indent_hanging_twips"]))

    # Alignment
    if atom.get("alignment"):
        jc = etree.SubElement(ppr, _w("jc"))
        jc.set(_w("val"), atom["alignment"])

    # rPr for paragraph mark (inherit style; just provide bold-on-mark if needed)
    ppr_rpr = etree.SubElement(ppr, _w("rPr"))

    # Build runs from segments (new style) OR text (old style)
    if "segments" in entry:
        for seg in entry["segments"]:
            text = seg.get("text", "")
            if not text:
                continue
            run = etree.SubElement(p, _w("r"))
            rpr = etree.SubElement(run, _w("rPr"))
            if seg.get("bold"):
                b = etree.SubElement(rpr, _w("b"))
                b.set(_w("val"), "true")
                bcs = etree.SubElement(rpr, _w("bCs"))
                bcs.set(_w("val"), "true")
            if seg.get("italic"):
                i = etree.SubElement(rpr, _w("i"))
                i.set(_w("val"), "true")
            if seg.get("underline"):
                u = etree.SubElement(rpr, _w("u"))
                u.set(_w("val"), "single")
            t = etree.SubElement(run, _w("t"))
            # Preserve leading/trailing whitespace
            if text != text.strip():
                t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t.text = text
    elif "text" in entry:
        run = etree.SubElement(p, _w("r"))
        rpr = etree.SubElement(run, _w("rPr"))
        if entry.get("bold"):
            b = etree.SubElement(rpr, _w("b"))
            b.set(_w("val"), "true")
            bcs = etree.SubElement(rpr, _w("bCs"))
            bcs.set(_w("val"), "true")
        if entry.get("italic"):
            i = etree.SubElement(rpr, _w("i"))
            i.set(_w("val"), "true")
        if entry.get("underline"):
            u = etree.SubElement(rpr, _w("u"))
            u.set(_w("val"), "single")
        t = etree.SubElement(run, _w("t"))
        text = entry["text"]
        if text != text.strip():
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t.text = text

    return p


# ---------- Public helpers used by main.py --------------------------------


def parse_template_from_b64(content_b64: str) -> Dict[str, Any]:
    """Convenience wrapper: take base64 (with or without data: prefix) and parse."""
    if "," in content_b64:
        content_b64 = content_b64.split(",", 1)[1]
    raw = base64.b64decode(content_b64)
    return parse_template(raw)
